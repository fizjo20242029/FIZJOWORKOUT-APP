import json
import os
import random
import io
import openpyxl
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
import streamlit as st
from groq import Groq


# ==============================================================================
# KONFIGURACJA STRONY STREAMLIT
# ==============================================================================
st.set_page_config(page_title="Fizjo Workout Ultimate", page_icon="💪", layout="centered")

if 'wylosowany_plan_cache' not in st.session_state:
    st.session_state.wylosowany_plan_cache = []
if 'historia_wiadomosci' not in st.session_state:
    st.session_state.historia_wiadomosci = [
        {"role": "system", "content": "Jesteś wirtualnym asystentem w aplikacji dla fizjoterapeutów i trenerów 'Fizjo Workout Ultimate'."}
    ]

# ==============================================================================
# BAZA FIZJOTERAPEUTYCZNA
# ==============================================================================
BAZA_FIZJO = {
    "Oddechowe": [
        {"nazwa": "Oddech przeponowy w leżeniu", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Leżenie tyłem, kolana ugięte. Jedna ręka na klatce, druga na brzuchu. Wdech nosem kieruje powietrze do brzucha (brzuch rośnie), wydech ustami.", "czas_min": 2, "parametry": "2-3 minuty", "miesnie": "Przepona, mięśnie międzyżebrowe"},
        {"nazwa": "Oddech metodą 'pursed lips'", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Wdech nosem, a następnie powolny, maksymalnie wydłużony wydech przez lekko przymknięte usta (jak przy delikatnym dmuchaniu na świeczkę).", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Przepona, mięśnie tłoczni brzusznej"},
        {"nazwa": "Oddech dolnożebrowy z taśmą", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Owiń taśmę elastyczną wokół dolnych żeber. Podczas wdechu staraj się rozepchnąć taśmę na boki, podczas wydechu taśma lekko uciska żebra.", "czas_min": 2, "parametry": "10-12 powtórzeń", "miesnie": "Mięśnie międzyżebrowe zewnętrzne, przepona"},
        {"nazwa": "Oddech asymetryczny na boku", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Leżenie na boku nieobjętym procesem chorobowym. Podczas wdechu kieruj powietrze do boku leżącego wyżej, rozciągając przestrzenie międzyżebrowe.", "czas_min": 2, "parametry": "5-8 powtórzeń na stronę", "miesnie": "Mięsień czworoboczny lędźwi, międzyżebrowe"},
        {"nazwa": "Pozycja krzesełkowa z oporem", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Siedząc, tułów lekko pochylony w przód, oparcie na przedramionach. Wydłużony wydech z oporem (np. dmuchanie przez rurkę do wody).", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Mięśnie pomocnicze wydechowe, przepona"},
        {"nazwa": "Wdech ze wznosem ramion", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Dynamiczny wdech nosem połączony z uniesieniem ramion przodem w górę, spokojny wydech ustami z opuszczaniem ramion luźno w dół.", "czas_min": 2, "parametry": "12 powtórzeń", "miesnie": "Mięsień piersiowy większy, zębaty przedni"},
        {"nazwa": "Oddech pudełkowy (Box Breathing)", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Wdech przez 4 sekundy, zatrzymanie powietrza na 4 sekundy, wydech przez 4 sekundy, zatrzymanie na bezdechu na 4 sekundy.", "czas_min": 2, "parametry": "4 pełne cykle", "miesnie": "Przepona, stabilizatory głębokie tułowia"},
        {"nazwa": "Mobilizacja klatki w rotacji", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Siad klęczny, jedna ręka za głową. Wykonuj powolny skręt tułowia w bok z głębokim wdechem w maksymalnym zakresie, powrót z wydechem.", "czas_min": 2, "parametry": "8-10 powtórzeń", "miesnie": "Mięśnie skośne brzucha, prostownik grzbietu piersiowy"},
        {"nazwa": "Oddech jednostronny w pozycji Sfinksa", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Leżenie przodem na przedramionach. Głęboki wdech z intencją skierowania powietrza do tylnych i dolnych partii płuc.", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Przepona, prostownik grzbietu"},
        {"nazwa": "Głęboki wdech w pozycji embrionalnej", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Ukłon japoński. Głębokie wdechy rozszerzające tylną ścianę klatki piersiowej.", "czas_min": 3, "parametry": "3 minuty", "miesnie": "Mięśnie międzyżebrowe tylne, przepona"},
        {"nazwa": "Oddychanie z oporem warg (Pursed lip breathing)", "opis": "INSTRUKCJA: Głęboki wdech nosem, a następnie bardzo powolny wydech przez zasznurowane usta (jak przy gwizdaniu), wytwarzając dodatnie ciśnienie w drogach oddechowych.", "czas_min": 2, "parametry": "10 powtórzeń", "miesnie": "Układ oddechowy (zapobieganie zapadaniu pęcherzyków)"},
        {"nazwa": "Naprzemienne oddychanie nozdrzami (Nadi Shodhana)", "opis": "INSTRUKCJA: Zatkaj prawe nozdrze, wdech lewym. Zatkaj lewe nozdrze, wydech prawym. Wdech prawym, wydech lewym. Reguluje układ autonomiczny.", "czas_min": 3, "parametry": "3 minuty", "miesnie": "Relaksacja układu nerwowego"},
    ],
    "Głowa/Szyja": [
        {"nazwa": "Retrakcja szyi (Cofanie brody)", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: W pozycji siedzącej cofnij głowę w płaszczyźnie poziomej (zrób 'podwójny podbródek'). Wzrok prosto, trzymaj 3 s.", "czas_min": 2, "parametry": "3 serie x 10 powtórzeń", "miesnie": "Mięśnie głębokie zginacze szyi, mięsień płatowaty głowy"},
        {"nazwa": "Izometryczne parcie w przód", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Przyłóż dłoń do czoła. Naciskaj głową w przód na stawiającą opór dłoń, nie pozwalając na ruch. Trzymaj przez 5 sekund.", "czas_min": 2, "parametry": "3 serie x 5s trzymania", "miesnie": "MOS, zginacze długie"},
        {"nazwa": "Izometryczne parcie w bok", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Przyłóż dłoń nad uchem. Naciskaj głową w bok przeciwko oporowi ręki. Trzymaj 5 sekund, zmień strony.", "czas_min": 2, "parametry": "3 serie x 5s na stronę", "miesnie": "Mięśnie pochyłe, mięsień płatowaty szyi"},
        {"nazwa": "Izometryczne parcie w tył", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Spleć dłonie z tyłu głowy. Naciskaj tyłem głowy w dłonie, utrzymując wzrok prosto. Utrzymaj 5 sekund.", "czas_min": 2, "parametry": "3 serie x 5s trzymania", "miesnie": "Mięsień podpotyliczny, prostownik grzbietu szyjny"},
        {"nazwa": "Rozciąganie trapeziusa", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Siedząc, opuść jedno ramię w dół. Drugą ręką delikatnie przyciągaj głowę DO przeciwnego barku.", "czas_min": 2, "parametry": "2 serie x 30 sekund", "miesnie": "Czworoboczny (góra), dźwigacz łopatki"},
        {"nazwa": "Rozciąganie dźwigacza łopatki", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Skręć głowę o 45 stopni w bok, a następnie skieruj brodę w dół do pachy. Delikatnie wspomóż ruch ręką.", "czas_min": 2, "parametry": "2 serie x 30 sekund na stronę", "miesnie": "Dźwigacz łopatki"},
        {"nazwa": "Otwieranie ust z oporem", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Podeprzyj żuchwę od dołu palcami. Otwieraj usta powoli, pokonując delikatny opór.", "czas_min": 2, "parametry": "3 serie x 10 powtórzeń", "miesnie": "Mięśnie nadgnykowe"},
        {"nazwa": "Ruchy sakadyczne oczu", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Trzymaj głowę nieruchomo. Przenoś wzrok szybko między dwoma punktami.", "czas_min": 1, "parametry": "1 minuta", "miesnie": "Mięśnie gałkoruchowe"},
        {"nazwa": "Automobilizacja z wałkiem", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Leżenie tyłem, twardy wałek pod podstawą czaszki. Wykonuj małe ruchy potakiwania.", "czas_min": 3, "parametry": "3 minuty", "miesnie": "Podpotyliczne"},
        {"nazwa": "Retrakcja z rotacją w leżeniu", "opis": "INSTRUKCJA RUCHU DLA PACJENTA: Leżenie przodem, czoło nad podłogą. Wykonuj retrakcję szyi, a następnie delikatny obrót głowy.", "czas_min": 2, "parametry": "10 powtórzeń na stronę", "miesnie": "Długi szyi"},
        {"nazwa": "PIR mięśnia czworobocznego (część zstępująca)", "opis": "INSTRUKCJA: Skłon głowy w bok do oporu. Delikatnie naciskaj głową na stawiającą opór dłoń (20% siły) przez 8s. Rozluźnij, zrób wydech i pogłęb zakres ruchu.", "czas_min": 3, "parametry": "3-5 powtórzeń na stronę", "miesnie": "Mięsień czworoboczny (góra)"},
        {"nazwa": "PIR mięśnia dźwigacza łopatki", "opis": "INSTRUKCJA: Głowa zgięta o 45 stopni i pochylona w stronę pachy. Zastosuj technikę napięcie-rozluźnienie-pogłębienie (jak w klasycznym PIR).", "czas_min": 3, "parametry": "3-5 powtórzeń na stronę", "miesnie": "Dźwigacz łopatki"},
        {"nazwa": "Autotrakcja szyi na ręczniku", "opis": "INSTRUKCJA: Owiń zrolowany ręcznik wokół potylicy, chwyć końce dłońmi. Pociągnij ręcznik do przodu i lekko w górę, wykonując delikatny wyprost szyi na stworzonym punkcie podparcia.", "czas_min": 2, "parametry": "10 powtórzeń", "miesnie": "Przestrzenie międzykręgowe szyjne"},
        {"nazwa": "Aktywacja głębokich zginaczy (Janda)", "opis": "INSTRUKCJA: Leżenie tyłem, głowa na płasko. Wykonaj sam ruch potakiwania (przyciągnięcie brody do krtani) bez odrywania tyłu głowy od kozetki.", "czas_min": 2, "parametry": "3x10 sekund", "miesnie": "Głębokie zginacze szyi (longus colli)"},
        {"nazwa": "Rozciąganie mięśnia mostkowo-obojczykowo-sutkowego (MOS)", "opis": "INSTRUKCJA: Jedną ręką ustabilizuj obojczyk. Odchyl głowę w tył, zrotuj w stronę przeciwną i wykonaj skłon do boku. Poczujesz ciągnięcie z przodu szyi.", "czas_min": 2, "parametry": "3x30 sekund na stronę", "miesnie": "Mięsień MOS, pochyłe"},
        {"nazwa": "Izometria szyi z piłką overball", "opis": "INSTRUKCJA: Umieść miękką piłkę między czołem a ścianą. Naciskaj na nią głową, utrzymując izometryczne napięcie mięśni bez ruchu w stawach.", "czas_min": 2, "parametry": "3x10 sekund (w każdą stronę)", "miesnie": "Globalne stabilizatory szyi"},
        
    ],
    "Kończyna górna": [
        {"nazwa": "Wznosy ramion w pozycji Y", "opis": "INSTRUKCJA: Leżenie przodem, ramiona pod kątem 45 stopni (litera Y), kciuki w górę. Unoś ramiona.", "czas_min": 3, "parametry": "3 serie x 12 powtórzeń", "miesnie": "Czworoboczny dolny"},
        {"nazwa": "Rotacja zewnętrzna z gumą", "opis": "INSTRUKCJA: Łokcie ugięte 90 stopni przy bokach. Rozciągaj gumę na boki poprzez rotację zewnętrzną w barkach.", "czas_min": 2, "parametry": "3 serie x 15 powtórzeń", "miesnie": "Stożek rotatorów"},
        {"nazwa": "Pompki plus (Scapular)", "opis": "INSTRUKCJA: W podporze przodem wypychaj łopatki w przód bez uginania łokci.", "czas_min": 2, "parametry": "3 serie x 12 powtórzeń", "miesnie": "Zębaty przedni"},
        {"nazwa": "Rozciąganie zginaczy nadgarstka", "opis": "INSTRUKCJA: Wyprostuj rękę, dłoń w górę. Drugą ręką chwyć palce i przyciągaj do siebie.", "czas_min": 2, "parametry": "2 serie x 30 sekund", "miesnie": "Zginacze nadgarstka"},
        {"nazwa": "Rozciąganie prostowników nadgarstka", "opis": "INSTRUKCJA: Wyprostuj rękę, dłoń w dół. Dociśnij dłoń w dół, zginając nadgarstek.", "czas_min": 2, "parametry": "2 serie x 30 sekund", "miesnie": "Prostowniki nadgarstka"},
        {"nazwa": "Odwodzenie hantli w opadzie", "opis": "INSTRUKCJA: Opad tułowia. Unoś ramiona na boki do wysokości tułowia.", "czas_min": 3, "parametry": "3 serie x 10 powtórzeń", "miesnie": "Naramienny tylny"},
        {"nazwa": "Ślizganie ramion po ścianie", "opis": "INSTRUKCJA: Oprzyj plecy o ścianę. Ręce ugięte 90 stopni. Przesuwaj ramiona w górę i w dół.", "czas_min": 2, "parametry": "3 serie x 12 powtórzeń", "miesnie": "Czworoboczny dolny"},
        {"nazwa": "Uginanie przedramion z supinacją", "opis": "INSTRUKCJA: W staniu, hantle chwytem neutralnym. Uginaj łokcie rotując dłonie.", "czas_min": 2, "parametry": "3 serie x 15 powtórzeń", "miesnie": "Biceps"},
        {"nazwa": "Prostowanie przedramion w opadzie", "opis": "INSTRUKCJA: Opad tułowia, łokcie 90 stopni. Prostuj ramię w tył.", "czas_min": 3, "parametry": "3 serie x 12 powtórzeń", "miesnie": "Triceps"},
        {"nazwa": "Spacer farmera z kettlebell", "opis": "INSTRUKCJA: Chwyć odważniki. Idź powolnym krokiem ze stabilnymi barkami.", "czas_min": 3, "parametry": "3 serie x 45 sekund", "miesnie": "Przedramiona, góra pleców"},
        {"nazwa": "Ślizg nerwu pośrodkowego", "opis": "INSTRUKCJA: Stojąc przy ścianie, ramię odwiedzione do 90st. Wykonuj powolny wyprost łokcia i nadgarstka, jednocześnie pochylając głowę w stronę przeciwną.", "czas_min": 2, "parametry": "15 powtórzeń na stronę", "miesnie": "Układ nerwowy obręczy barkowej"},
        {"nazwa": "Drenaż limfatyczny kończyny (Automasaż)", "opis": "INSTRUKCJA: Wykonuj delikatne, głaskające ruchy od palców dłoni w kierunku węzłów chłonnych pod pachą. Ruch ma przesunąć skórę, nie naciskać mocno na mięśnie.", "czas_min": 5, "parametry": "5 minut na stronę", "miesnie": "Układ limfatyczny"},
        {"nazwa": "Ćwiczenia czynno-bierne kończyny", "opis": "INSTRUKCJA: Używając zdrowej ręki, powoli i płynnie unoś, zginaj i prostuj osłabioną rękę w bezbolesnym zakresie ruchu.", "czas_min": 3, "parametry": "10 powtórzeń na każdy staw", "miesnie": "Stawy kończyny górnej (Neurologia)"},
        {"nazwa": "Wahadła Codmana (Zwis luźny)", "opis": "INSTRUKCJA: Opad tułowia, zdrowa ręka podparta o stół. Chora ręka zwisa luźno. Ruchy kółkowe ramienia generuj poprzez balansowanie tułowiem, bark jest całkowicie wiotki.", "czas_min": 3, "parametry": "2 minuty", "miesnie": "Rozluźnienie torebki stawowej barku"},
        {"nazwa": "Sleeper Stretch (Rozciąganie torebki tylnej)", "opis": "INSTRUKCJA: Leżenie na chorym boku, ramię ugięte 90st przed sobą. Zdrową ręką powoli dociskaj przedramię do leżanki (rotacja wewnętrzna) do uczucia ciągnięcia.", "czas_min": 2, "parametry": "3x30 sekund", "miesnie": "Torebka tylna stawu ramiennego, podgrzebieniowy"},
        {"nazwa": "Izometria rotacji z piłką o ścianę", "opis": "INSTRUKCJA: Stojąc bokiem do ściany, łokieć ugięty 90st. Przyciśnij małą piłkę do ściany grzbietem dłoni (rotacja zew.) i pchaj z 30% siły bez ruchu.", "czas_min": 2, "parametry": "3x8 sekund", "miesnie": "Stożek rotatorów"},
        {"nazwa": "Wzorzec PNF (D1 Zgięcie) z taśmą", "opis": "INSTRUKCJA: Stań na gumie. Chwyć ją po przekątnej i prowadź ramię od wyprostu/rotacji wewn. (przy przeciwległym biodrze) do pełnego zgięcia/rotacji zewn. nad głową.", "czas_min": 3, "parametry": "3x10 powtórzeń", "miesnie": "Łańcuch kinematyczny kończyny górnej"},
        {"nazwa": "Automasaż rozcięgna dłoniowego piłeczką", "opis": "INSTRUKCJA: Oprzyj dłoń na twardej piłeczce (np. golfowej/lacrosse). Wykonuj powolne, dogniatające ruchy rolowania po całej wewnętrznej powierzchni dłoni.", "czas_min": 2, "parametry": "2 minuty na stronę", "miesnie": "Rozcięgno dłoniowe, zginacze palców"},
        {"nazwa": "Ślizg nerwu promieniowego", "opis": "INSTRUKCJA: Stojąc, opuść bark. Ramię wyprostowane, zrotowane do wewnątrz (kciuk w stronę uda), nadgarstek zgięty dłoniowo. Pochylaj głowę w stronę przeciwną odciągając ramię w tył.", "czas_min": 2, "parametry": "15 powtórzeń na stronę", "miesnie": "Nerw promieniowy"},
        {"nazwa": "Ślizg nerwu łokciowego", "opis": "INSTRUKCJA: 'Maska Batmana'. Ramię odwiedzione do 90st. Zegnij łokieć i nadgarstek, próbując przyłożyć dłoń do oka/ucha jak maskę. Jednocześnie pochyl głowę w przeciwną stronę.", "czas_min": 2, "parametry": "15 powtórzeń na stronę", "miesnie": "Nerw łokciowy"},
        {"nazwa": "Rozciąganie w futrynie drzwi (Piersiowy)", "opis": "INSTRUKCJA: Oprzyj przedramiona na futrynie drzwi (łokcie na wys. barków lub wyżej). Wykonaj powolny wykrok jedną nogą w przód, rozciągając klatkę piersiową.", "czas_min": 2, "parametry": "3x30 sekund", "miesnie": "Piersiowy większy i mniejszy"},
        {"nazwa": "PIR mięśnia piersiowego większego", "opis": "INSTRUKCJA: Leżenie tyłem na brzegu kozetki, ręka zwisa. Lekko naciskaj ręką w górę przeciw oporowi (20% siły) przez 8s. Rozluźnij z wydechem i pozwól ręce opaść niżej.", "czas_min": 3, "parametry": "4 cykle na stronę", "miesnie": "Piersiowy większy"},
    
    ],
    "Core (Tułów)": [
        {"nazwa": "Plank (Podpór przodem)", "opis": "INSTRUKCJA: Oprzyj się na przedramionach i palcach stóp. Ciało w linii prostej.", "czas_min": 3, "parametry": "3 serie x 30 sekund", "miesnie": "Core"},
        {"nazwa": "Side Plank (Podpór bokiem)", "opis": "INSTRUKCJA: Leżenie na boku, podparcie na przedramieniu. Unieś biodra.", "czas_min": 3, "parametry": "2 serie x 20 sekund na stronę", "miesnie": "Skośne brzucha"},
        {"nazwa": "Dead Bug (Martwy robak)", "opis": "INSTRUKCJA: Leżenie tyłem. Opuszczaj przeciwną rękę i nogę, lędźwie w matę.", "czas_min": 3, "parametry": "3 serie x 10 powtórzeń na stronę", "miesnie": "Poprzeczny brzucha"},
        {"nazwa": "Bird Dog (Pies-ptak)", "opis": "INSTRUKCJA: Klęk podparty. Wyciągnij prawą rękę w przód i lewą nogę w tył.", "czas_min": 3, "parametry": "3 serie x 10 powtórzeń na stronę", "miesnie": "Prostownik grzbietu"},
        {"nazwa": "Hollow Body (Pozycja kołyski)", "opis": "INSTRUKCJA: Leżenie tyłem. Oderwij łopatki i proste nogi od podłogi.", "czas_min": 2, "parametry": "3 serie x 20 sekund", "miesnie": "Prosty brzucha"},
        {"nazwa": "Bear Crawl Hold (Niedźwiedź)", "opis": "INSTRUKCJA: Klęk podparty. Unieś kolana 2 cm nad ziemię.", "czas_min": 2, "parametry": "3 serie x 20 sekund", "miesnie": "Poprzeczny brzucha"},
        {"nazwa": "Niedźwiedź dynamiczny", "opis": "INSTRUKCJA: Kroki w podporze niedźwiedzia.", "czas_min": 3, "parametry": "3 serie x 30 sekund", "miesnie": "Core"},
        {"nazwa": "Opuszczanie obustronne nóg", "opis": "INSTRUKCJA: Leżenie tyłem. Powoli opuszczaj obie proste nogi.", "czas_min": 2, "parametry": "3 serie x 10 powtórzeń", "miesnie": "Dół brzucha"},
        {"nazwa": "Skośne spięcia brzucha", "opis": "INSTRUKCJA: Leżenie tyłem. Kieruj ramię w stronę przeciwnego kolana.", "czas_min": 2, "parametry": "3 serie x 15 powtórzeń", "miesnie": "Skośne brzucha"},
        {"nazwa": "Russian Twist", "opis": "INSTRUKCJA: Siad z uniesionymi stopami. Rotuj klatkę piersiową.", "czas_min": 2, "parametry": "3 serie x 20 skrętów", "miesnie": "Skośne brzucha"},
        {"nazwa": "Przeprosty McKenziego (Leżenie przodem)", "opis": "INSTRUKCJA: Leżenie przodem. Dłonie na wysokości barków. Powoli prostuj łokcie, unosząc samą klatkę piersiową, miednica zostaje przyklejona do maty.", "czas_min": 2, "parametry": "3x10 powtórzeń", "miesnie": "Prostownik grzbietu"},
        {"nazwa": "Trening mięśni dna miednicy (Kegla)", "opis": "INSTRUKCJA: W wygodnej pozycji napinaj mięśnie dna miednicy (jak przy próbie zatrzymania strumienia moczu). Utrzymaj napięcie, oddychając swobodnie.", "czas_min": 3, "parametry": "10 spięć po 5 sekund", "miesnie": "Mięśnie dna miednicy"},
        {"nazwa": "Zegary miednicy (Pelvic tilts)", "opis": "INSTRUKCJA: Leżenie tyłem, kolana ugięte. Płynnie wciskaj odcinek lędźwiowy w matę (tyłopochylenie), a następnie odrywaj go, robiąc łuk (przodopochylenie).", "czas_min": 2, "parametry": "20 płynnych powtórzeń", "miesnie": "Dolny odcinek lędźwiowy, miednica"},
        {"nazwa": "Ślizg nerwu udowego", "opis": "INSTRUKCJA: Leżenie przodem. Zginaj kolano (pięta do pośladka) przy jednoczesnym wyproście szyi (patrz przed siebie). Następnie prostuj kolano i zginaj szyję w dół.", "czas_min": 2, "parametry": "15 powtórzeń na stronę", "miesnie": "Nerw udowy"},
        {"nazwa": "Pozycja odciążająca (Trakcja 90/90)", "opis": "INSTRUKCJA: Leżenie tyłem, podudzia oparte na krześle lub dużej piłce (kąty 90st w biodrach i kolanach). Oddychaj głęboko do brzucha.", "czas_min": 5, "parametry": "5-10 minut", "miesnie": "Odbarczenie krążków międzykręgowych L-S"},
        {"nazwa": "PIR mięśnia czworobocznego lędźwi (QL)", "opis": "INSTRUKCJA: Leżenie na boku zdrowym na wałku, chora noga zwisa poza kozetkę. Izometryczne uniesienie miednicy w górę (5s), rozluźnienie i opadnięcie nogi grawitacyjnie w dół.", "czas_min": 3, "parametry": "4 cykle", "miesnie": "Czworoboczny lędźwi (QL)"},
        {"nazwa": "Rolowanie klatki piersiowej na wałku", "opis": "INSTRUKCJA: Leżenie tyłem na wałku (wałek w poprzek kręgosłupa piersiowego). Ręce za głową, wykonuj delikatne wyprosty odcinka piersiowego przez wałek.", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Prostownik grzbietu piersiowy, stawy międzykręgowe"},
        {"nazwa": "Mobilizacja przepony (automasaż)", "opis": "INSTRUKCJA: W leżeniu tyłem, zegnij kolana. Wsuń delikatnie opuszki palców pod dolne żebra. Na wydechu delikatnie masuj i rozluźniaj tkanki pod łukiem żebrowym.", "czas_min": 3, "parametry": "3 minuty", "miesnie": "Przepona"},
        {"nazwa": "Rotacje lędźwiowe w leżeniu tyłem", "opis": "INSTRUKCJA: Leżenie tyłem, ręce szeroko na boki. Kolana zgięte i złączone. Powoli opuszczaj złączone kolana raz na lewą, raz na prawą stronę, nie odrywając łopatek.", "czas_min": 3, "parametry": "20 powtórzeń", "miesnie": "Mięśnie skośne brzucha, rotatory lędźwiowe"},
        {"nazwa": "Pozycja ukłonu japońskiego (Child's pose)", "opis": "INSTRUKCJA: Siad na piętach. Wyciągnij ramiona daleko w przód po macie, czoło oprzyj na podłodze. Rozluźnij całkowicie dolny odcinek pleców.", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Prostownik grzbietu (rozciąganie)"},
        {"nazwa": "Wzmacnianie wielodzielnego (Cofanie miednicy)", "opis": "INSTRUKCJA: Klęk podparty. Wykonaj bardzo delikatny ruch cofnięcia miednicy (tyłopochylenie) wykorzystując głębokie mięśnie przykręgosłupowe. Ruch jest mikroskopijny.", "czas_min": 2, "parametry": "3x15 powtórzeń", "miesnie": "Mięsień wielodzielny (multifidus)"},
        
    ],
    "Kończyna dolna": [
        {"nazwa": "Przysiad klasyczny (Squat)", "opis": "INSTRUKCJA: Stopy na szerokość barków. Schodź biodrami w dół i w tył.", "czas_min": 2, "parametry": "3 serie x 12 powtórzeń", "miesnie": "Czworogłowy uda, pośladki"},
        {"nazwa": "Mostki biodrowe (Glute Bridge)", "opis": "INSTRUKCJA: Leżenie tyłem, kolana ugięte. Unoś biodra w górę.", "czas_min": 2, "parametry": "3 serie x 15 powtórzeń", "miesnie": "Pośladkowy wielki"},
        {"nazwa": "Wznosy na palce stojąc", "opis": "INSTRUKCJA: Unoś pięty maksymalnie w górę, przechodząc na palce.", "czas_min": 2, "parametry": "3 serie x 20 powtórzeń", "miesnie": "Łydki"},
        {"nazwa": "Zakroki (Reverse Lunges)", "opis": "INSTRUKCJA: Duży krok w tył, opuszczając biodra w dół.", "czas_min": 3, "parametry": "3 serie x 10 powtórzeń na nogę", "miesnie": "Czworogłowy uda"},
        {"nazwa": "Odwodzenie nogi z gumą", "opis": "INSTRUKCJA: Guma nad kostkami. Odprowadzaj nogę w bok.", "czas_min": 2, "parametry": "3 serie x 12 na stronę", "miesnie": "Pośladkowy średni"},
        {"nazwa": "Przysiad bułgarski", "opis": "INSTRUKCJA: Jedna stopa oparta z tyłu. Przysiad na nodze wykrocznej.", "czas_min": 3, "parametry": "3 serie x 8 na nogę", "miesnie": "Czworogłowy uda"},
        {"nazwa": "Clamshell (Muszelka)", "opis": "INSTRUKCJA: Leżenie na boku, kolana ugięte. Unieś górne kolano.", "czas_min": 2, "parametry": "3 serie x 15 na stronę", "miesnie": "Pośladkowy średni"},
        {"nazwa": "Unoszenie prostej nogi (SLR)", "opis": "INSTRUKCJA: Unoś prostą nogę do wysokości drugiego kolana.", "czas_min": 2, "parametry": "3 serie x 12 na nogę", "miesnie": "Prosty uda"},
        {"nazwa": "Martwy ciąg na jednej nodze", "opis": "INSTRUKCJA: Stojąc na jednej nodze, wykonaj skłon w przód.", "czas_min": 3, "parametry": "3 serie x 8 na stronę", "miesnie": "Dwugłowy uda"},
        {"nazwa": "Krzesełko przy ścianie", "opis": "INSTRUKCJA: Oprzyj plecy o ścianę, zejdź biodrami do kąta 90 stopni. Trzymaj nieruchomo.", "czas_min": 2, "parametry": "3x30 sekund", "miesnie": "Czworogłowy uda, stabilizacja kolana"},
        {"nazwa": "Ślizg nerwu kulszowego (Neuromobilizacja)", "opis": "INSTRUKCJA: Leżenie tyłem, noga ugięta w biodrze do 90st. Prostuj kolano zadzierając palce stopy na siebie (napięcie), a opuszczając stopę obciągaj palce (luzowanie).", "czas_min": 2, "parametry": "15 powtórzeń na stronę", "miesnie": "Układ nerwowy, tylna taśma"},
        {"nazwa": "Rozciąganie zginaczy biodra w klęku", "opis": "INSTRUKCJA: Klęk jednonóż. Wypchnij miednicę mocno do przodu, zachowując wyprostowany tułów, aż poczujesz rozciąganie z przodu uda nogi zakrocznej.", "czas_min": 2, "parametry": "3x30 sekund na stronę", "miesnie": "Mięsień biodrowo-lędźwiowy, prosty uda"},
        {"nazwa": "Ćwiczenia równoważne na jednej nodze (Propriocepcja)", "opis": "INSTRUKCJA: Stań na jednej nodze, lekko zginając kolano. Staraj się utrzymać równowagę. Dla utrudnienia zamknij oczy lub stań na poduszce.", "czas_min": 2, "parametry": "3x30 sekund na nogę", "miesnie": "Stabilizatory stawu skokowego i kolanowego"},
        {"nazwa": "TKE (Terminal Knee Extension) z gumą", "opis": "INSTRUKCJA: Guma zaczepiona o drabinkę i tył Twojego kolana. Zaczynasz od ugiętego kolana, następnie wpychasz je mocno w tył do pełnego wyprostu, napinając udo.", "czas_min": 3, "parametry": "3x15 powtórzeń", "miesnie": "Głowa przyśrodkowa (VMO) czworogłowego"},
        {"nazwa": "Protokół Alfredsona (Ekscentryka Achillesa)", "opis": "INSTRUKCJA: Stojąc na krawędzi stopnia, wejdź na palce obunóż. Przenieś ciężar na jedną nogę i BARDZO POWOLI opuszczaj piętę w dół. Wróć w górę pomagając sobie drugą nogą.", "czas_min": 3, "parametry": "3x15 powtórzeń na nogę", "miesnie": "Ścięgno Achillesa, brzuchaty łydki"},
        {"nazwa": "Krótka stopa (Short foot exercise)", "opis": "INSTRUKCJA: Stojąc boso, staraj się zbliżyć głowę pierwszej kości śródstopia do pięty, podnosząc łuk stopy bez zaginania palców (palce leżą płasko).", "czas_min": 2, "parametry": "3x10 powtórzeń", "miesnie": "Mięśnie wewnętrzne stopy"},
        {"nazwa": "Ślizgi piętą po leżance (Heel slides)", "opis": "INSTRUKCJA: Leżenie tyłem (np. po zabiegu na kolano). Powoli przyciągaj piętę do pośladka, ślizgając nią po podłożu, aż do granicy bólu, następnie powoli wyprostuj.", "czas_min": 3, "parametry": "3x15 powtórzeń", "miesnie": "Zginacze stawu kolanowego, poprawa ROM"},
        {"nazwa": "PIR grupy kulszowo-goleniowej", "opis": "INSTRUKCJA: Leżenie tyłem, noga uniesiona do oporu. Lekko dociśnij nogę w dół przeciw oporowi (np. taśmy) przez 8s. Rozluźnij i zwiększ zgięcie w biodrze.", "czas_min": 3, "parametry": "4 cykle na nogę", "miesnie": "Grupa kulszowo-goleniowa"},
        {"nazwa": "PIR mięśnia gruszkowatego", "opis": "INSTRUKCJA: Leżenie tyłem. Chora noga ugięta, stopa za kolanem nogi zdrowej. Przyciągaj kolano po przekątnej do przeciwnego barku z techniką izometrii.", "czas_min": 2, "parametry": "4 cykle na stronę", "miesnie": "Mięsień gruszkowaty"},
        {"nazwa": "Izometryczne ściskanie piłki (VMO)", "opis": "INSTRUKCJA: Siad z nogami wyprostowanymi, mała piłka między kolanami. Ściskaj piłkę kolanami i utrzymuj napięcie przez 5 sekund.", "czas_min": 2, "parametry": "3x10 powtórzeń", "miesnie": "Przywodziciele, głowa przyśrodkowa (VMO)"},
        {"nazwa": "Nordic Hamstring Curl (Ekscentryka)", "opis": "INSTRUKCJA: Klęk, terapeuta (lub drabinka) blokuje stopy. Bardzo powoli opadaj tułowiem w przód, hamując ruch wyłącznie tyłem ud, aż oprzesz się na rękach.", "czas_min": 3, "parametry": "3x5 powtórzeń", "miesnie": "Grupa kulszowo-goleniowa (ekscentryka)"},
        {"nazwa": "Zbieranie ręcznika palcami stóp", "opis": "INSTRUKCJA: W siadzie. Ułóż ręcznik na podłodze, postaw na nim stopy. Chwytaj palcami stóp materiał i podwijaj pod siebie.", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Zginacze palców, rozcięgno podeszwowe"},
        {"nazwa": "Hip Airplanes (Samoloty)", "opis": "INSTRUKCJA: W opadzie tułowia na jednej nodze (pozycja jaskółki). Wykonuj rotację miednicy na zewnątrz (otwieranie biodra) i do wewnątrz (zamykanie).", "czas_min": 3, "parametry": "3x8 powtórzeń na nogę", "miesnie": "Rotatory biodra, stabilizatory miednicy"},
        {"nazwa": "Ślizg nerwu strzałkowego", "opis": "INSTRUKCJA: W leżeniu lub siadzie, noga ugięta. Wykonaj zgięcie podeszwowe stopy (obciągnięcie palców) z jednoczesną inwersją stopy do wewnątrz, połączone ze zgięciem szyi.", "czas_min": 2, "parametry": "15 powtórzeń na stronę", "miesnie": "Nerw strzałkowy"},
        {"nazwa": "PIR mięśnia prostego uda", "opis": "INSTRUKCJA: Leżenie przodem. Zegnij kolano. Naciskaj stopą w dół na rękę stawiającą opór (5s). Rozluźnij i dociągnij piętę bliżej pośladka.", "czas_min": 3, "parametry": "4 cykle na nogę", "miesnie": "Prosty uda"},
        {"nazwa": "Rolowanie pasma biodrowo-piszczelowego (ITB)", "opis": "INSTRUKCJA: Podpór bokiem na przedramieniu z wałkiem pod udem. Powoli roluj zewnętrzną część uda od biodra aż nad kolano. Unikaj rolowania samych stawów.", "czas_min": 3, "parametry": "2 minuty na stronę", "miesnie": "Pasmo biodrowo-piszczelowe (ITB), TFL"},
        {"nazwa": "Korekcja miednicy (Drop pelvis)", "opis": "INSTRUKCJA: Stań jedną nogą na stepie/schodku, druga noga zwisa. Kontrolowanie opuszczaj zwisającą nogę (biodro opada) i unoś z powrotem aktywując pośladek nogi podporowej.", "czas_min": 3, "parametry": "3x12 powtórzeń na nogę", "miesnie": "Pośladkowy średni, stabilizacja miednicy"},
        
    ],

    "Neurologia i Koordynacja": [
        {"nazwa": "Ćwiczenia Frenkla (Ślizganie piętą)", "opis": "INSTRUKCJA: Leżenie tyłem. Ze wzrokową kontrolą ruchu powoli przesuwaj piętę jednej nogi wzdłuż piszczeli drugiej nogi, od kostki aż do kolana i z powrotem.", "czas_min": 3, "parametry": "10 powtórzeń na nogę", "miesnie": "Koordynacja, propriocepcja, układ nerwowy"},
        {"nazwa": "Przetaczanie tułowia (wzorce rotacyjne)", "opis": "INSTRUKCJA: Leżenie tyłem. Zainicjuj obrót na bok poprzez sięgnięcie samą ręką (góra ciała) w bok, miednica podąża za ruchem z opóźnieniem.", "czas_min": 3, "parametry": "10 przetoczeń na stronę", "miesnie": "Taśmy skośne, integracja odruchów"},
        {"nazwa": "PNF - Wzorzec miednicy (Elewacja przednia)", "opis": "INSTRUKCJA: Leżenie na boku. Pociągnij miednicę w górę i w przód (kierunek do pępka) pokonując opór terapeuty lub taśmy.", "czas_min": 2, "parametry": "3x10 powtórzeń", "miesnie": "Mięśnie tułowia, skośne brzucha, QL"},
        {"nazwa": "Trening dwuzadaniowy (Dual-tasking)", "opis": "INSTRUKCJA: Wykonuj proste ćwiczenie ruchowe (np. maszerowanie w miejscu lub odbijanie balonu) z jednoczesnym zadaniem poznawczym (np. liczenie w tył od 100 co 7).", "czas_min": 5, "parametry": "5 minut", "miesnie": "Neuroplastyczność mózgu"}
    ]
}

# ==============================================================================
# BAZA TRENINGOWA NA SIŁOWNI
# ==============================================================================
BAZA_SILOWNIA = {
    "Rozgrzewka": [
        {"nazwa": "Monster Walk z mini-bandem", "opis": "INSTRUKCJA: Załóż mini-band nad kolana, przyjmij pozycję półprzysiadu. Wykonuj szerokie kroki w bok, pilnując kolan na zewnątrz.", "czas_min": 2, "parametry": "2 min", "miesnie": "Pośladkowy średni, stabilizatory"},
        {"nazwa": "Face Pulls na wyciągu", "opis": "INSTRUKCJA: Przyciągaj linę do twarzy, rozchylając dłonie na boki i mocno spinając tył barków.", "czas_min": 2, "parametry": "3x15", "miesnie": "Naramienne tył, czworoboczny"},
        {"nazwa": "Rotacje zewnętrzne z gumą", "opis": "INSTRUKCJA: Łokieć przy boku pod kątem 90 stopni, rotuj przedramię na zewnątrz.", "czas_min": 2, "parametry": "3x12", "miesnie": "Stożek rotatorów"},
        {"nazwa": "Koci grzbiet", "opis": "INSTRUKCJA: W klęku podpartym naprzemiennie wyginaj kręgosłup w górę i w dół.", "czas_min": 2, "parametry": "10 razy", "miesnie": "Prostowniki grzbietu"},
        {"nazwa": "Skip A w miejscu", "opis": "INSTRUKCJA: Dynamiczne unoszenie kolan z zachowaniem wyprostowanej sylwetki.", "czas_min": 2, "parametry": "2 min", "miesnie": "Biodra, core"},
        {"nazwa": "Krążenia ramion z gumami", "opis": "INSTRUKCJA: Obszerne krążenia z lekką taśmą w dłoniach.", "czas_min": 2, "parametry": "20 razy", "miesnie": "Barki"},
        {"nazwa": "Pajacyki", "opis": "INSTRUKCJA: Klasyczne skoki z odwodzeniem ramion i nóg.", "czas_min": 2, "parametry": "2 min", "miesnie": "Całe ciało"},
        {"nazwa": "Wymachy nóg w przód", "opis": "INSTRUKCJA: Stojąc, dynamicznie wymachuj prostą nogą w przód.", "czas_min": 2, "parametry": "10 na stronę", "miesnie": "Biodra, dwugłowe"},
        {"nazwa": "Przysiady z masą ciała", "opis": "INSTRUKCJA: Płynne przysiady bez obciążenia, pilnując ustawienia stóp.", "czas_min": 2, "parametry": "20 razy", "miesnie": "Nogi"},
        {"nazwa": "Pompki przy ścianie", "opis": "INSTRUKCJA: Wstępna aktywacja klatki przez odpychanie się od ściany.", "czas_min": 2, "parametry": "15 razy", "miesnie": "Klatka, triceps"},
        {"nazwa": "Wykroki z rotacją", "opis": "INSTRUKCJA: Wykonaj wykrok i skręt tułowia w stronę nogi wykrocznej.", "czas_min": 3, "parametry": "5 na stronę", "miesnie": "Biodra, kręgosłup"},
        {"nazwa": "Skakanka", "opis": "INSTRUKCJA: Skoki przez skakankę w stałym tempie.", "czas_min": 3, "parametry": "3 min", "miesnie": "Łydki, serce"},
        {"nazwa": "Plank", "opis": "INSTRUKCJA: Podpór przodem na przedramionach, ciało w linii prostej.", "czas_min": 2, "parametry": "1 min", "miesnie": "Core"},
        {"nazwa": "Bird-Dog", "opis": "INSTRUKCJA: W klęku podpartym unoś przeciwną rękę i nogę.", "czas_min": 2, "parametry": "10 na stronę", "miesnie": "Core"},
        {"nazwa": "Wiosłowanie taśmą", "opis": "INSTRUKCJA: Zaczep taśmę o słup i przyciągaj końce do bioder.", "czas_min": 2, "parametry": "20 razy", "miesnie": "Plecy"},
        {"nazwa": "Mostki biodrowe", "opis": "INSTRUKCJA: Leżenie tyłem, unoś biodra poprzez napięcie pośladków.", "czas_min": 2, "parametry": "20 razy", "miesnie": "Pośladki"},
        {"nazwa": "Spacer w podporze", "opis": "INSTRUKCJA: Z pozycji stojącej przejdź dłońmi do podporu i wróć.", "czas_min": 2, "parametry": "10 razy", "miesnie": "Core, ramiona"},
        {"nazwa": "Krążenia bioder", "opis": "INSTRUKCJA: Obszerne ruchy miednicą w staniu.", "czas_min": 1, "parametry": "1 min", "miesnie": "Biodra"},
        {"nazwa": "Rowerek - lekki opór", "opis": "INSTRUKCJA: Spokojne pedałowanie w celu rozgrzania stawów.", "czas_min": 3, "parametry": "3 min", "miesnie": "Nogi"},
        {"nazwa": "Wspinaczka górska (Mountain Climbers)", "opis": "INSTRUKCJA: W podporze dynamicznie przyciągaj kolana do klatki.", "czas_min": 2, "parametry": "1 min", "miesnie": "Core, nogi"}
    ],
    "Klatka piersiowa": [
        {"nazwa": "Wyciskanie sztangi płasko", "opis": "POZYCJA: Połóż się na ławce, stopy stabilnie na podłożu, łopatki ściągnięte. RUCH: Opuść sztangę do dolnej części klatki, a następnie wyciśnij dynamicznie w górę. UWAGA: Nie odrywaj pośladków od ławki, kontroluj ciężar przy opuszczaniu.", "czas_min": 3, "parametry": "4x8", "miesnie": "Klatka, triceps"},
        {"nazwa": "Wyciskanie hantli skos dodatni", "opis": "POZYCJA: Ławka ustawiona pod kątem 30 stopni. Stopy stabilne. RUCH: Wyciskaj hantle nad klatkę, zbliżając je do siebie w górnej fazie. UWAGA: Nie przeprostowuj łokci w szczycie ruchu.", "czas_min": 3, "parametry": "4x10", "miesnie": "Góra klatki, barki"},
        {"nazwa": "Dipy (pompki na poręczach)", "opis": "POZYCJA: Podeprzyj się na poręczach, tułów pochylony w przód. RUCH: Uginaj ramiona, aż poczujesz rozciągnięcie klatki, następnie wróć do pełnego wyprostu. UWAGA: Prowadź łokcie szerzej niż przy tricepsie.", "czas_min": 3, "parametry": "3x10", "miesnie": "Klatka, triceps"},
        {"nazwa": "Rozpiętki z hantlami", "opis": "POZYCJA: Leżąc płasko, hantle nad klatką, lekko ugięte łokcie. RUCH: Otwieraj ramiona szeroko na boki, aż poczujesz mocne rozciągnięcie, wróć do pozycji początkowej. UWAGA: Ruch wykonuj w stawie barkowym, nie zmieniaj kąta ugięcia łokci.", "czas_min": 2, "parametry": "3x12", "miesnie": "Klatka"},
        {"nazwa": "Wyciskanie hantli płasko", "opis": "POZYCJA: Leżenie na ławce płaskiej. RUCH: Wyciskaj hantle w pełnym zakresie ruchu, dbając o symetrię. UWAGA: Kontroluj tor ruchu hantli, aby nie traciły stabilności.", "czas_min": 3, "parametry": "4x10", "miesnie": "Klatka"},
        {"nazwa": "Pompki klasyczne", "opis": "POZYCJA: Podpora przodem, dłonie szerzej niż barki. RUCH: Obniż tułów do ziemi, zachowując linię prostą, wróć do góry. UWAGA: Napinaj brzuch, aby uniknąć zapadania się lędźwi.", "czas_min": 2, "parametry": "3xMax", "miesnie": "Klatka, triceps, core"},
        {"nazwa": "Brama (Cable Crossover)", "opis": "POZYCJA: Stań w wykroku w środku bramy. RUCH: Ściągaj uchwyty wyciągu w dół i do środka, spinając klatkę. UWAGA: Prowadź ruch po łuku, nie wyginaj kręgosłupa.", "czas_min": 3, "parametry": "3x15", "miesnie": "Klatka"},
        {"nazwa": "Wyciskanie Hammer Strength", "opis": "POZYCJA: Siedząc na maszynie, chwyć uchwyty. RUCH: Wypychaj uchwyty przed siebie. UWAGA: Nie odrywaj pleców od oparcia podczas wysiłku.", "czas_min": 3, "parametry": "3x12", "miesnie": "Klatka, triceps"},
        {"nazwa": "Pompki na podwyższeniu nóg", "opis": "POZYCJA: Stopy na ławce, dłonie na ziemi. RUCH: Pompka klasyczna. UWAGA: Skupienie na górnej części klatki, pilnuj bioder.", "czas_min": 2, "parametry": "3x12", "miesnie": "Góra klatki, barki"},
        {"nazwa": "Wyciskanie w maszynie Smitha", "opis": "POZYCJA: Leżenie na ławce pod gryfem maszyny. RUCH: Wyciskaj gryf, który porusza się po szynie. UWAGA: Dopasuj pozycję ławki, aby gryf trafiał na linię mostka.", "czas_min": 3, "parametry": "4x10", "miesnie": "Klatka"},
        {"nazwa": "Floor Press (wyciskanie z podłogi)", "opis": "POZYCJA: Leżenie na plecach, sztanga nad klatką. RUCH: Opuść sztangę, aż tricepsy dotkną podłogi, następnie wyciśnij. UWAGA: Zatrzymuj ruch na podłodze, nie odbijaj.", "czas_min": 3, "parametry": "3x10", "miesnie": "Klatka, triceps"},
        {"nazwa": "Pompki diamentowe", "opis": "POZYCJA: Podpór przodem, dłonie blisko siebie (kciuki i palce wskazujące stykają się). RUCH: Pompka. UWAGA: Skupienie na środku klatki i pracy tricepsów.", "czas_min": 2, "parametry": "3x10", "miesnie": "Środek klatki, triceps"},
        {"nazwa": "Landmine Press (jednorącz)", "opis": "POZYCJA: Stojąc, koniec sztangi zamocowany w rogu. RUCH: Wyciskaj koniec sztangi jedną ręką przed siebie. UWAGA: Stabilizuj tułów napięciem brzucha.", "czas_min": 3, "parametry": "3x12", "miesnie": "Góra klatki, barki"},
        {"nazwa": "Wyciskanie hantla jednorącz leżąc", "opis": "POZYCJA: Leżąc, jedna ręka pracuje, druga stabilizuje. RUCH: Wyciskaj hantel w górę. UWAGA: Utrzymuj stabilną pozycję barków na ławce.", "czas_min": 3, "parametry": "3x12", "miesnie": "Klatka"},
        {"nazwa": "Maszyna Butterfly (motylki)", "opis": "POZYCJA: Siedząc, przedramiona na poduszkach maszyny. RUCH: Zbliżaj ramiona do siebie. UWAGA: Nie przesuwaj ramion za wysoko, łokcie na linii barków.", "czas_min": 2, "parametry": "3x15", "miesnie": "Klatka"},
        {"nazwa": "Wyciskanie sztangi skos ujemny", "opis": "POZYCJA: Ławka ustawiona pod skosem w dół. RUCH: Wyciskaj sztangę. UWAGA: Kontroluj tempo, aby sztanga nie spadła na szyję.", "czas_min": 3, "parametry": "3x10", "miesnie": "Dół klatki"},
        {"nazwa": "Pompki z taśmą oporową", "opis": "POZYCJA: Taśma przełożona przez plecy, dłonie trzymają końce. RUCH: Pompka z dodatkowym oporem. UWAGA: Pilnuj, by taśma nie ześlizgnęła się z pleców.", "czas_min": 2, "parametry": "3x10", "miesnie": "Klatka"},
        {"nazwa": "Pullover z hantlem", "opis": "POZYCJA: Leżąc prostopadle do ławki. RUCH: Przenoś hantel zza głowy nad klatkę po łuku. UWAGA: Rozciągaj klatkę w bezpiecznym zakresie.", "czas_min": 2, "parametry": "3x12", "miesnie": "Klatka, najszerszy"},
        {"nazwa": "Wyciskanie hantli chwytem neutralnym", "opis": "POZYCJA: Leżenie płasko. RUCH: Wyciskaj hantle, trzymając dłonie skierowane do siebie. UWAGA: Bardzo stabilna praca dla barków i klatki.", "czas_min": 3, "parametry": "3x10", "miesnie": "Klatka"},
        {"nazwa": "Wyciskanie sztangi wąsko", "opis": "POZYCJA: Chwyt sztangi na szerokość barków. RUCH: Wyciskaj, prowadząc łokcie blisko tułowia. UWAGA: Mocna praca tricepsa, ale też środek klatki.", "czas_min": 3, "parametry": "3x8", "miesnie": "Środek klatki, triceps"}
    ],
    "Plecy": [
        {"nazwa": "Podciąganie na drążku nachwytem", "opis": "POZYCJA: Chwyć drążek szerzej niż barki, ramiona wyprostowane. RUCH: Podciągnij klatkę piersiową do drążka, prowadząc łokcie w dół i do tyłu. UWAGA: Unikaj huśtania ciałem; ruch powinien być kontrolowany w obie strony.", "czas_min": 3, "parametry": "3 serie x Max", "miesnie": "Najszerszy grzbietu, obły większy"},
        {"nazwa": "Wiosłowanie sztangą w opadzie", "opis": "POZYCJA: Stopy na szerokość bioder, kolana lekko ugięte. Pochyl tułów do kąta 45 stopni, plecy proste. RUCH: Przyciągaj sztangę do dolnej części brzucha, ściągając łopatki. UWAGA: Kręgosłup musi być w linii prostej przez cały czas.", "czas_min": 3, "parametry": "4 serie x 10 powtórzeń", "miesnie": "Plecy, prostownik grzbietu, biceps"},
        {"nazwa": "Martwy ciąg klasyczny", "opis": "POZYCJA: Stopy na szerokość bioder, sztanga blisko piszczeli. RUCH: Podnieś sztangę, prostując biodra i kolana jednocześnie, wypychając klatkę w przód. UWAGA: Nie wyginaj kręgosłupa w łuk; ciężar prowadź jak najbliżej nóg.", "czas_min": 4, "parametry": "4 serie x 6 powtórzeń", "miesnie": "Prostowniki grzbietu, pośladki, plecy"},
        {"nazwa": "Wiosłowanie hantlem jednorącz", "opis": "POZYCJA: Jedna ręka i kolano na ławce, plecy równolegle do podłoża. RUCH: Przyciągaj hantel w stronę biodra, łokieć prowadź blisko tułowia. UWAGA: Unikaj rotacji tułowia podczas ruchu.", "czas_min": 3, "parametry": "3 serie x 12 na stronę", "miesnie": "Najszerszy grzbietu, czworoboczny"},
        {"nazwa": "Ściąganie drążka wyciągu górnego", "opis": "POZYCJA: Usiądź, zablokuj uda pod wałkiem. RUCH: Przyciągaj drążek do górnej części klatki piersiowej, lekko odchylając się w tył. UWAGA: Nie bujaj ciałem, pracuj świadomie mięśniami pleców.", "czas_min": 2, "parametry": "3 serie x 12 powtórzeń", "miesnie": "Najszerszy grzbietu"},
        {"nazwa": "Wiosłowanie na maszynie (Hammer)", "opis": "POZYCJA: Siedząc na maszynie, chwyć uchwyty. RUCH: Przyciągaj uchwyty do brzucha, mocno spinając łopatki w szczycie. UWAGA: Nie odrywaj klatki od oparcia maszyny.", "czas_min": 3, "parametry": "3 serie x 12 powtórzeń", "miesnie": "Mięśnie pleców"},
        {"nazwa": "Face Pulls z linką", "opis": "POZYCJA: Stań przed wyciągiem górnym. RUCH: Przyciągaj linkę do czoła, rozchylając dłonie na boki i rotując ramiona na zewnątrz. UWAGA: Łokcie muszą znajdować się wyżej niż dłonie.", "czas_min": 2, "parametry": "3 serie x 15 powtórzeń", "miesnie": "Czworoboczny, tył barków"},
        {"nazwa": "Wiosłowanie na pasach TRX", "opis": "POZYCJA: Chwyć uchwyty, odchyl ciało w tył w pozycji deski. RUCH: Przyciągaj klatkę do dłoni, pracując łopatkami. UWAGA: Napnij pośladki i brzuch, aby ciało nie zwisało.", "czas_min": 2, "parametry": "3 serie x 15 powtórzeń", "miesnie": "Plecy, core"},
        {"nazwa": "Pullover hantlem (leżąc)", "opis": "POZYCJA: Leżenie w poprzek ławki, głowa poza nią. RUCH: Przenoś hantel zza głowy nad klatkę po łuku. UWAGA: Ruch wykonuj powoli, pilnując stabilności obręczy barkowej.", "czas_min": 2, "parametry": "3 serie x 12 powtórzeń", "miesnie": "Najszerszy grzbietu, klatka"},
        {"nazwa": "Wiosłowanie kettlem", "opis": "POZYCJA: Szeroki rozkrok, kettlebell na ziemi między stopami. RUCH: Podciągaj kettlebell do brzucha w opadzie. UWAGA: Plecy muszą pozostać proste, nie zaokrąglaj lędźwi.", "czas_min": 3, "parametry": "3 serie x 12 powtórzeń", "miesnie": "Plecy"},
        {"nazwa": "Ściąganie uchwytu V na wyciągu", "opis": "POZYCJA: Siedząc, chwyt wąski. RUCH: Przyciągaj uchwyt V do splotu słonecznego. UWAGA: Prowadź ruch płynnie, bez szarpania ciężarem.", "czas_min": 2, "parametry": "3 serie x 12 powtórzeń", "miesnie": "Najszerszy grzbietu"},
        {"nazwa": "Shrugsy (wzruszanie barkami)", "opis": "POZYCJA: Stojąc, hantle w dłoniach wzdłuż tułowia. RUCH: Wzruszaj barkami w górę, jakbyś chciał dotknąć nimi uszu. UWAGA: Nie rotuj barkami, ruch tylko góra-dół.", "czas_min": 2, "parametry": "4 serie x 12 powtórzeń", "miesnie": "Czworoboczny"},
        {"nazwa": "Wiosłowanie taśmą oporową", "opis": "POZYCJA: Taśma zaczepiona o stabilny słup. RUCH: Przyciągaj końce taśmy do bioder, pilnując postawy. UWAGA: Kontroluj fazę powrotu (odpuszczania) taśmy.", "czas_min": 2, "parametry": "3 serie x 20 powtórzeń", "miesnie": "Plecy"},
        {"nazwa": "Podciąganie podchwytem", "opis": "POZYCJA: Chwyt drążka od dołu. RUCH: Podciągaj klatkę w górę, mocno angażując mięśnie ramion i pleców. UWAGA: Pełen zakres ruchu w dół.", "czas_min": 3, "parametry": "3 serie x Max", "miesnie": "Plecy, biceps"},
        {"nazwa": "Dzień dobry (Good Mornings)", "opis": "POZYCJA: Sztanga na karku, stopy na szerokość bioder. RUCH: Skłon tułowia z prostymi (lekko ugiętymi) nogami, biodra w tył. UWAGA: Ruch tylko do momentu, gdy poczujesz rozciąganie dwugłowych.", "czas_min": 3, "parametry": "3 serie x 10 powtórzeń", "miesnie": "Prostowniki grzbietu, dwugłowe"},
        {"nazwa": "Wiosłowanie na ławce skosnej", "opis": "POZYCJA: Leżenie klatką na ławce skośnej. RUCH: Wiosłuj hantlami do boków ciała. UWAGA: Nie odrywaj klatki od ławeczki.", "czas_min": 3, "parametry": "3 serie x 12 powtórzeń", "miesnie": "Plecy, tył barków"},
        {"nazwa": "High Pulls (sztanga)", "opis": "POZYCJA: Sztanga w dłoniach, stopy na szerokość bioder. RUCH: Dynamiczne wyprostowanie bioder z jednoczesnym podciągnięciem sztangi do brody. UWAGA: Użyj siły nóg do zainicjowania ruchu.", "czas_min": 3, "parametry": "3 serie x 8 powtórzeń", "miesnie": "Plecy, barki"},
        {"nazwa": "Wiosłowanie na wyciągu dolnym", "opis": "POZYCJA: Siedząc z wyprostowanymi nogami. RUCH: Przyciągaj rączkę do brzucha, wypinając klatkę w przód. UWAGA: Unikaj odchylania się nadmiernie w tył.", "czas_min": 2, "parametry": "3 serie x 12 powtórzeń", "miesnie": "Plecy"},
        {"nazwa": "T-Bar Row (wiosłowanie T-Bar)", "opis": "POZYCJA: W opadzie, chwyt za uchwyty maszyny. RUCH: Przyciągaj ciężar do klatki. UWAGA: Pilnuj, by plecy były płaskie, nie zaokrąglaj się.", "czas_min": 3, "parametry": "4 serie x 10 powtórzeń", "miesnie": "Plecy"},
        {"nazwa": "Spacer farmera", "opis": "POZYCJA: Wyprostowana sylwetka, hantle w dłoniach. RUCH: Marsz krótkimi krokami, stabilizując tułów. UWAGA: Nie pozwól, aby ciężar ciągnął ramiona w dół, trzymaj barki aktywnie.", "czas_min": 3, "parametry": "3 serie x 30m", "miesnie": "Plecy, chwyt, core"}
    ],

    "Barki": [
        {"nazwa": "Wyciskanie żołnierskie (OHP)", "opis": "POZYCJA: Stojąc, sztanga na wysokości obojczyków, pośladki i brzuch mocno napięte. RUCH: Wyciśnij sztangę nad głowę, lekko cofając głowę. UWAGA: Nie wyginaj lędźwi w łuk, trzymaj sztywny core.", "czas_min": 3, "parametry": "4x8", "miesnie": "Naramienny przedni, triceps, core"},
        {"nazwa": "Wyciskanie hantli siedząc", "opis": "POZYCJA: Siedząc na ławce (oparcie 75-80 stopni), hantle na wysokości uszu. RUCH: Wyciskaj hantle w górę, nie blokując łokci w szczycie. UWAGA: Prowadź łokcie lekko przed linią tułowia.", "czas_min": 3, "parametry": "4x10", "miesnie": "Naramienny przedni i boczny"},
        {"nazwa": "Wznosy hantli bokiem", "opis": "POZYCJA: Stojąc, hantle wzdłuż ciała, łokcie minimalnie ugięte. RUCH: Unoś ramiona na boki do wysokości barków. UWAGA: Mały palec dłoni powinien być wyżej niż kciuk (jak przy wylewaniu wody z dzbanka).", "czas_min": 2, "parametry": "4x15", "miesnie": "Naramienny boczny"},
        {"nazwa": "Wznosy hantli w opadzie tułowia", "opis": "POZYCJA: Opad tułowia, plecy proste, hantle zwisają swobodnie. RUCH: Unoś ramiona szeroko na boki, spinając tył barków. UWAGA: Nie używaj pędu, ruch ma być izolowany.", "czas_min": 2, "parametry": "4x15", "miesnie": "Naramienny tylny, czworoboczny"},
        {"nazwa": "Face Pulls z liną wyciągu", "opis": "POZYCJA: Stojąc, wyciąg ustawiony na wysokości twarzy. RUCH: Przyciągaj linę do czoła, mocno rotując ramiona na zewnątrz w końcowej fazie. UWAGA: Skup się na ściągnięciu łopatek i napięciu tyłu barków.", "czas_min": 2, "parametry": "3x15", "miesnie": "Naramienny tylny, stożek rotatorów"},
        {"nazwa": "Wyciskanie Arnolda", "opis": "POZYCJA: Siedząc, hantle przed twarzą, wnętrza dłoni skierowane do siebie. RUCH: Wyciskaj hantle w górę, jednocześnie rotując nadgarstki na zewnątrz. UWAGA: Płynny ruch łączący pracę przodu i boku barku.", "czas_min": 3, "parametry": "3x12", "miesnie": "Naramienne (całość)"},
        {"nazwa": "Podciąganie sztangi wzdłuż tułowia", "opis": "POZYCJA: Stojąc, szeroki chwyt sztangi. RUCH: Podciągaj sztangę do linii klatki piersiowej, prowadząc łokcie wysoko w górę. UWAGA: Zbyt wąski chwyt może obciążać stawy barkowe.", "czas_min": 3, "parametry": "3x12", "miesnie": "Naramienny boczny, czworoboczny"},
        {"nazwa": "Wznosy bokiem na wyciągu jednorącz", "opis": "POZYCJA: Stojąc bokiem do wyciągu dolnego, linka trzymana w przeciwnej ręce (za plecami lub przed). RUCH: Unoś ramię w bok. UWAGA: Linka zapewnia stałe napięcie mięśnia przez cały ruch.", "czas_min": 2, "parametry": "3x15", "miesnie": "Naramienny boczny"},
        {"nazwa": "Szrugsy ze sztangą / hantlami", "opis": "POZYCJA: Stojąc prosto, ciężar w opuszczonych rękach. RUCH: Wykonuj wzruszenia barkami w górę. UWAGA: Tylko ruch pionowy (góra-dół), nie wykonuj krążeń barkami.", "czas_min": 2, "parametry": "4x15", "miesnie": "Czworoboczny (kaptury)"},
        {"nazwa": "Odwrotne rozpiętki na maszynie (Reverse Pec Deck)", "opis": "POZYCJA: Siedząc przodem do maszyny, ręce wyprostowane przed sobą. RUCH: Odwodź ramiona maksymalnie w tył. UWAGA: Skup się na pracy tyłu barku, a nie całych pleców.", "czas_min": 2, "parametry": "3x15", "miesnie": "Naramienny tylny"},
        {"nazwa": "Wyciskanie na maszynie siedząc (Shoulder Press)", "opis": "POZYCJA: Siedząc na maszynie, plecy mocno dociśnięte. RUCH: Wypychaj uchwyty w górę. UWAGA: Świetne ćwiczenie do bezpiecznego obciążenia barków bez ryzyka utraty równowagi.", "czas_min": 3, "parametry": "3x12", "miesnie": "Naramienny przedni i boczny"},
        {"nazwa": "Unoszenie hantli w przód", "opis": "POZYCJA: Stojąc prosto. RUCH: Unoś hantle naprzemiennie przed siebie do linii wzroku. UWAGA: Nie odchylaj tułowia w tył.", "czas_min": 2, "parametry": "3x12 na rękę", "miesnie": "Naramienny przedni"},
        {"nazwa": "Unoszenie talerza w przód (z rotacją)", "opis": "POZYCJA: Stojąc, trzymając talerz oburącz (jak kierownicę). RUCH: Unoś talerz przed siebie i lekko skręcaj. UWAGA: Maksymalne spięcie przodu barków.", "czas_min": 2, "parametry": "3x15", "miesnie": "Naramienny przedni"},
        {"nazwa": "Push Press (Wybicie sztangi)", "opis": "POZYCJA: Jak do OHP. RUCH: Lekko ugnij kolana, a następnie wybij dynamicznie sztangę w górę pomagając sobie nogami. UWAGA: Ćwiczenie budujące dynamikę i siłę.", "czas_min": 3, "parametry": "4x6", "miesnie": "Całe barki, triceps"},
        {"nazwa": "Wyciskanie hantli jednorącz stojąc", "opis": "POZYCJA: Stojąc, hantel w jednej ręce, druga dla przeciwwagi. RUCH: Wyciskaj ciężar nad głowę. UWAGA: Nymaga bardzo mocnego spięcia mięśni core, aby nie przechylić tułowia.", "czas_min": 3, "parametry": "3x10 na stronę", "miesnie": "Naramienny przedni, core"},
        {"nazwa": "Wznosy bokiem w siadzie", "opis": "POZYCJA: Siedząc na ławce. RUCH: Wznosy hantli na boki. UWAGA: Pozycja siedząca całkowicie eliminuje możliwość oszukiwania z nóg i tułowia.", "czas_min": 2, "parametry": "3x15", "miesnie": "Naramienny boczny"},
        {"nazwa": "Wznosy hantli leżąc przodem na ławce skośnej", "opis": "POZYCJA: Klatka oparta o ławkę skośną (około 30-45 stopni). RUCH: Odwodź ręce w tył jak przy wiosłowaniu szerokim. UWAGA: Idealna izolacja tyłu barku bez obciążania lędźwi.", "czas_min": 2, "parametry": "3x15", "miesnie": "Naramienny tylny"},
        {"nazwa": "Lu Raises (Pełne wznosy z hantlami)", "opis": "POZYCJA: Stojąc prosto. RUCH: Unoś powoli hantle z boku ciała aż nad samą głowę po pełnym łuku. UWAGA: Poprawia mobilność obręczy barkowej.", "czas_min": 3, "parametry": "3x10", "miesnie": "Naramienny boczny, trapezius"},
        {"nazwa": "Odwrotne rozpiętki z gumą oporową", "opis": "POZYCJA: Stojąc, trzymaj gumę przed sobą na wys. klatki. RUCH: Rozciągaj gumę na boki do momentu aż dotknie klatki. UWAGA: Trzymaj łokcie wysoko.", "czas_min": 2, "parametry": "3x20", "miesnie": "Naramienny tylny"},
        {"nazwa": "Y-Raises leżąc przodem (na ławce)", "opis": "POZYCJA: Klatka oparta na ławce skośnej. RUCH: Unoś hantle w kształt litery 'Y' (pod kątem 45 stopni przed siebie). UWAGA: Kciuki skierowane w sufit.", "czas_min": 2, "parametry": "3x12", "miesnie": "Dolny czworoboczny, barki"}
    ],
    "Brzuch / Core": [
        {"nazwa": "Allahy (Skłony na wyciągu górnym)", "opis": "POZYCJA: Klęcząc tyłem lub przodem do wyciągu, chwyć linę za głową. RUCH: Zwijaj tułów (jak przy brzuszkach), przyciągając klatkę do miednicy. UWAGA: Ruch ma zachodzić w kręgosłupie, to nie jest skłon w biodrach.", "czas_min": 2, "parametry": "4x15", "miesnie": "Prosty brzucha"},
        {"nazwa": "Ab Wheel (Kółko rehabilitacyjne)", "opis": "POZYCJA: Klęk podparty, dłonie na kółku. RUCH: Wyjeżdżaj kółkiem w przód, napinając mocno brzuch, i wracaj. UWAGA: Nie pozwól, aby lędźwia zapadły się w dół (hiperlordoza).", "czas_min": 3, "parametry": "3x10", "miesnie": "Prosty brzucha, core"},
        {"nazwa": "Wznosy nóg w zwisie na drążku", "opis": "POZYCJA: Zwis swobodny na drążku. RUCH: Unoś proste nogi (lub ugięte kolana) w stronę klatki piersiowej. UWAGA: Staraj się podwinąć miednicę w górnej fazie ruchu, unikaj huśtania.", "czas_min": 2, "parametry": "3xMax", "miesnie": "Dół brzucha, biodrowo-lędźwiowy"},
        {"nazwa": "Plank (Deska) z obciążeniem", "opis": "POZYCJA: Podpór na przedramionach, talerz ułożony na plecach. RUCH: Utrzymuj sztywną linię ciała. UWAGA: Napnij mocno pośladki i wciągnij pępek.", "czas_min": 2, "parametry": "3x45 sekund", "miesnie": "Core (stabilizacja globalna)"},
        {"nazwa": "Russian Twist z obciążeniem", "opis": "POZYCJA: Siad, nogi uniesione, kettlebell lub talerz w dłoniach. RUCH: Skręcaj tułów, przenosząc ciężar z jednej strony na drugą. UWAGA: Wzrok i klatka piersiowa podążają za ciężarem.", "czas_min": 2, "parametry": "3x20 (na stronę)", "miesnie": "Mięśnie skośne brzucha"},
        {"nazwa": "Woodchopper na wyciągu", "opis": "POZYCJA: Stojąc bokiem do wyciągu górnego. RUCH: Ciągnij uchwyt po przekątnej z góry na dół (jak przy rąbaniu drewna), rotując tułów. UWAGA: Trzymaj ręce względnie proste, pracuj rotacją tułowia.", "czas_min": 2, "parametry": "3x12 (na stronę)", "miesnie": "Skośne brzucha, core"},
        {"nazwa": "Dead Bug (Martwy robak) z hantlami", "opis": "POZYCJA: Leżenie tyłem, ręce i nogi w górze, lędźwia wciśnięte w matę. RUCH: Opuszczaj naprzemiennie przeciwną rękę i nogę. UWAGA: Jeśli lędźwia odrywają się od podłogi, nie opuszczaj nóg tak nisko.", "czas_min": 2, "parametry": "3x12 (na stronę)", "miesnie": "Core, poprzeczny brzucha"},
        {"nazwa": "Świeca (Dragon Flag - wersja uproszczona)", "opis": "POZYCJA: Leżenie tyłem na ławce, dłonie trzymają ławkę za głową. RUCH: Unoś spięte, proste ciało w górę, opierając się tylko na barkach, następnie powoli opuszczaj. UWAGA: Maksymalne napięcie całego ciała.", "czas_min": 3, "parametry": "3x8", "miesnie": "Prosty brzucha (całość)"},
        {"nazwa": "Spacer farmera jednorącz", "opis": "POZYCJA: Stojąc, ciężki hantel / kettlebell w jednej dłoni. RUCH: Marsz z zachowaniem idealnie pionowej sylwetki. UWAGA: Mięśnie core muszą mocno pracować, by zapobiec wygięciu tułowia w stronę ciężaru.", "czas_min": 3, "parametry": "3x40 metrów (na stronę)", "miesnie": "Skośne brzucha, czworoboczny lędźwi"},
        {"nazwa": "Pallof Press (Wyciskanie z gumą)", "opis": "POZYCJA: Stojąc bokiem do wyciągu lub gumy na wysokości klatki, uchwyt w dłoniach na mostku. RUCH: Wyciskaj uchwyt prosto przed siebie i przytrzymaj. UWAGA: Walcz z siłą rotacyjną, która chce Cię skręcić w stronę wyciągu.", "czas_min": 2, "parametry": "3x12 (na stronę)", "miesnie": "Core (antyrotacja)"},
        {"nazwa": "Brzuszki na ławce skośnej ujemnej", "opis": "POZYCJA: Nogi zablokowane na ławce ze skosem w dół. RUCH: Skłony tułowia. UWAGA: Zejdź w dół pod kontrolą i wracaj dynamiką (pełne rozwinięcie i zwinięcie mięśnia prostego).", "czas_min": 2, "parametry": "3x15", "miesnie": "Prosty brzucha"},
        {"nazwa": "Toes to Bar (Palce do drążka)", "opis": "POZYCJA: Zwis swobodny. RUCH: Z silnym napięciem dołu brzucha przyciągnij stopy aż do samego drążka. UWAGA: Unikaj 'kippingu' (bujania z bioder).", "czas_min": 3, "parametry": "3x10", "miesnie": "Cały brzuch, zginacze bioder"},
        {"nazwa": "V-Ups (Scyzoryki)", "opis": "POZYCJA: Leżenie płasko na macie. RUCH: Jednocześnie unieś z podłogi złączone, proste nogi i tułów, dotykając dłońmi stóp w górze. UWAGA: Ekstremalne napięcie centrum.", "czas_min": 2, "parametry": "3x15", "miesnie": "Prosty brzucha"},
        {"nazwa": "Side Plank Dips (Opuszczanie bioder w desce)", "opis": "POZYCJA: Deska boczna na przedramieniu. RUCH: Opuść biodro aż musnie podłogę, następnie wypchnij je dynamicznie wysoko w górę. UWAGA: Głowa i szyja w jednej linii z ciałem.", "czas_min": 2, "parametry": "3x15 na stronę", "miesnie": "Skośne brzucha"},
        {"nazwa": "Bicycle Crunches (Rowerek z rotacją)", "opis": "POZYCJA: Leżenie tyłem, ręce za głową. RUCH: Przyciągaj prawe kolano i lewy łokieć do siebie, jednocześnie prostując lewą nogę, potem zmiana. UWAGA: Wolne tempo daje najlepsze efekty.", "czas_min": 2, "parametry": "3x20", "miesnie": "Skośne i prosty brzucha"},
        {"nazwa": "Spacer z odważnikiem nad głową (Overhead Carry)", "opis": "POZYCJA: Kettlebell lub hantel utrzymany na wyprostowanej ręce idealnie nad głową. RUCH: Powolny, stabilny marsz. UWAGA: Mega praca dla stabilizacji core pod obciążeniem u góry.", "czas_min": 3, "parametry": "3x30 metrów na stronę", "miesnie": "Core, obręcz barkowa"},
        {"nazwa": "Skręty tułowia na maszynie (Torso Rotation)", "opis": "POZYCJA: W siadzie na specjalistycznej maszynie. RUCH: Ograniczona i ułożyskowana rotacja w boki. UWAGA: Bardzo łatwy sposób na wyizolowanie mięśni skośnych w kontrolowanych warunkach.", "czas_min": 2, "parametry": "3x15 na stronę", "miesnie": "Skośne brzucha"},
        {"nazwa": "Spięcia brzucha na maszynie (Ab Crunch Machine)", "opis": "POZYCJA: W siadzie, chwyt za górne uchwyty. RUCH: Kulenie całego ciała jak przy rolowaniu dywanu. UWAGA: Pamiętaj o wydechu w fazie największego skurczu.", "czas_min": 2, "parametry": "3x15", "miesnie": "Prosty brzucha"},
        {"nazwa": "Hollow Body Rock (Kołyska)", "opis": "POZYCJA: Leżenie z uniesionymi z ziemi ramionami za głową i nogami, lędźwie mocno w podłodze (kształt banana). RUCH: Lekkie kołysanie się przód-tył w pełnym napięciu. UWAGA: Lędźwia nie mogą odrywać się od ziemi.", "czas_min": 2, "parametry": "3x30 sekund", "miesnie": "Core (izometria)"},
        {"nazwa": "Przenoszenie nóg nad ławką w siadzie", "opis": "POZYCJA: Siad oparty lekko w tył, proste nogi. RUCH: Przenoszenie połączonych, wyprostowanych nóg z prawej na lewą stronę ponad np. ustawionym na sztorc hantlem. UWAGA: Utrzymuj tułów zablokowany.", "czas_min": 2, "parametry": "3x20 przeskoczeń", "miesnie": "Dół brzucha, biodra"}
    ],
    
    "Ręce": [
        {"nazwa": "Uginanie sztangi łamanej", "opis": "POZYCJA: Stojąc, chwyć sztangę podchwytem na szerokość barków. RUCH: Uginaj ramiona w łokciach, prowadząc gryf w stronę klatki piersiowej. UWAGA: Łokcie muszą pozostać w jednej pozycji przy tułowiu, nie wykonuj zamachu tułowiem.", "czas_min": 2, "parametry": "3x12", "miesnie": "Biceps, m. ramienny"},
        {"nazwa": "Uginanie hantli z supinacją", "opis": "POZYCJA: Stojąc z hantlami w dłoniach, chwyt neutralny. RUCH: Uginaj ramię, jednocześnie obracając nadgarstek na zewnątrz (dłoń do góry). UWAGA: Rotacja następuje w trakcie ruchu, nie na końcu.", "czas_min": 2, "parametry": "3x12", "miesnie": "Biceps"},
        {"nazwa": "Wyciskanie francuskie sztangi", "opis": "POZYCJA: Leżąc na ławce, sztanga nad czołem. RUCH: Zginaj łokcie, opuszczając sztangę do czoła, następnie prostuj ramiona. UWAGA: Łokcie nie mogą rozchodzić się na boki.", "czas_min": 2, "parametry": "3x10", "miesnie": "Triceps"},
        {"nazwa": "Prostowanie linek na wyciągu", "opis": "POZYCJA: Stojąc, łokcie przyklejone do talii. RUCH: Prostuj ramiona, rozciągając linkę na dole. UWAGA: Skup się na pełnym wyproście w łokciu bez ruchu barków.", "czas_min": 2, "parametry": "3x15", "miesnie": "Triceps"},
        {"nazwa": "Uginanie hantli młotkowe", "opis": "POZYCJA: Stojąc, hantle chwytem neutralnym (dłonie do siebie). RUCH: Uginaj ramię bez rotacji nadgarstka. UWAGA: Trzymaj nadgarstek sztywno przez cały czas.", "czas_min": 2, "parametry": "3x12", "miesnie": "Mięsień ramienny, biceps"},
        {"nazwa": "Dipy na ławeczce", "opis": "POZYCJA: Ręce na ławeczce za plecami, stopy na podłożu. RUCH: Obniżaj biodra, uginając łokcie do kąta prostego, potem wróć. UWAGA: Plecy muszą prowadzić blisko ławki.", "czas_min": 2, "parametry": "3x15", "miesnie": "Triceps"},
        {"nazwa": "Uginanie na modlitewniku", "opis": "POZYCJA: Przedramiona oparte o oparcie modlitewnika. RUCH: Uginaj ramię, pilnując, by łokieć nie odrywał się od oparcia. UWAGA: Nie wykonuj ruchu w pełnym wyproście, aby chronić stawy.", "czas_min": 2, "parametry": "3x12", "miesnie": "Biceps"},
        {"nazwa": "Prostowanie hantla za głowę", "opis": "POZYCJA: Siedząc lub stojąc, hantel oburącz nad głową. RUCH: Zginaj łokcie, opuszczając hantel za kark, potem prostuj. UWAGA: Łokcie powinny być skierowane do przodu/góry, nie na boki.", "czas_min": 2, "parametry": "3x12", "miesnie": "Triceps (głowa długa)"},
        {"nazwa": "Uginanie z gumą oporową", "opis": "POZYCJA: Stań na środku gumy, końce w dłoniach. RUCH: Uginaj ramiona z gumą. UWAGA: Kontroluj fazę powrotu (odpuszczania) gumy, nie pozwól, by szarpnęła ręką.", "czas_min": 2, "parametry": "3x20", "miesnie": "Biceps"},
        {"nazwa": "Zottman Curl", "opis": "POZYCJA: Stojąc z hantlami. RUCH: Ugnij ręce z supinacją (dłoń do góry), w szczycie obróć nadgarstki i opuszczaj hantle chwytem pronowanym (dłonie w dół). UWAGA: Płynna rotacja w szczycie.", "czas_min": 2, "parametry": "3x10", "miesnie": "Biceps, przedramiona"},
        {"nazwa": "Triceps podchwytem (wyciąg)", "opis": "POZYCJA: Stojąc przy wyciągu, chwyt jednorącz podchwytem. RUCH: Prostuj ramię w dół. UWAGA: Łokieć jest osią obrotu, trzymaj go nieruchomo przy ciele.", "czas_min": 2, "parametry": "3x12", "miesnie": "Triceps"},
        {"nazwa": "Spider Curl (na ławce skośnej)", "opis": "POZYCJA: Leżąc klatką na ławce skośnej, ramiona pionowo w dół. RUCH: Uginaj hantle w górę. UWAGA: Nie wykonuj zamachu łokciami, trzymaj je w jednej linii.", "czas_min": 2, "parametry": "3x12", "miesnie": "Biceps"},
        {"nazwa": "Pompki diamentowe", "opis": "POZYCJA: Dłonie ułożone w kształt diamentu. RUCH: Pompka. UWAGA: Podczas ruchu łokcie prowadź blisko tułowia.", "czas_min": 2, "parametry": "3x12", "miesnie": "Triceps"},
        {"nazwa": "Uginanie koncentryczne", "opis": "POZYCJA: W siadzie, łokieć oparty o wewnętrzną stronę kolana. RUCH: Uginaj ramię. UWAGA: Nie wspomagaj się drugą ręką.", "czas_min": 2, "parametry": "3x12", "miesnie": "Biceps"},
        {"nazwa": "Skull Crushers z hantlami", "opis": "POZYCJA: Leżenie, hantle nad głową. RUCH: Opuszczaj hantle po bokach głowy, prostuj. UWAGA: Stabilne nadgarstki.", "czas_min": 2, "parametry": "3x10", "miesnie": "Triceps"},
        {"nazwa": "Uginanie z kettlem", "opis": "POZYCJA: Chwyt za kulę kettla oburącz. RUCH: Uginaj ramiona do klatki. UWAGA: Stabilna pozycja kręgosłupa.", "czas_min": 2, "parametry": "3x12", "miesnie": "Biceps"},
        {"nazwa": "Prostowanie za głowę z gumą", "opis": "POZYCJA: Guma zaczepiona nisko, za plecami. RUCH: Prostuj ręce nad głowę. UWAGA: Utrzymuj stabilny korpus.", "czas_min": 2, "parametry": "3x15", "miesnie": "Triceps"},
        {"nazwa": "Uginanie nadgarstków sztangą", "opis": "POZYCJA: Przedramiona oparte o ławkę, dłonie poza nią. RUCH: Uginaj nadgarstki. UWAGA: Małe, kontrolowane ruchy.", "czas_min": 2, "parametry": "3x15", "miesnie": "Przedramiona"},
        {"nazwa": "System 21 (Biceps)", "opis": "POZYCJA: Stojąc. RUCH: 7 powtórzeń do połowy w górę, 7 od połowy do góry, 7 pełnych. UWAGA: Bardzo męczące, pilnuj techniki przy każdym odcinku.", "czas_min": 3, "parametry": "3x21", "miesnie": "Biceps"},
        {"nazwa": "Wyprosty ramion jednorącz", "opis": "POZYCJA: Stojąc przy wyciągu. RUCH: Prostuj ramię do pełnego wyprostu. UWAGA: Pełne wyczucie pracy tricepsa.", "czas_min": 2, "parametry": "3x15", "miesnie": "Triceps"}
    ],
    "Nogi": [
        {"nazwa": "Przysiad ze sztangą na plecach", "opis": "POZYCJA: Stopy na szerokość barków, sztanga na mięśniach czworobocznych (nie na kręgach). RUCH: Wypchnij biodra w tył i zginaj kolana, utrzymując proste plecy. UWAGA: Kolana powinny podążać za linią palców, nie schodź do środka.", "czas_min": 4, "parametry": "4x8", "miesnie": "Czworogłowe, pośladki, przywodziciele"},
        {"nazwa": "Martwy ciąg na prostych nogach", "opis": "POZYCJA: Sztanga w dłoniach, stopy na szerokość bioder, kolana lekko ugięte. RUCH: Wypychaj biodra mocno w tył, prowadząc sztangę blisko nóg do połowy piszczeli. UWAGA: Kręgosłup musi być idealnie prosty, nie garb się.", "czas_min": 3, "parametry": "4x10", "miesnie": "Dwugłowe uda, pośladki, prostowniki"},
        {"nazwa": "Wykroki z hantlami", "opis": "POZYCJA: Stojąc, hantle w dłoniach wzdłuż tułowia. RUCH: Wykonaj duży krok w przód, uginając oba kolana do kąta 90 stopni. UWAGA: Tułów trzymaj pionowo, nie pochylaj się do przodu.", "czas_min": 3, "parametry": "3x12 na nogę", "miesnie": "Czworogłowe, pośladki"},
        {"nazwa": "Wypychanie na suwnicy", "opis": "POZYCJA: Siedząc na maszynie, stopy na platformie. RUCH: Wypychaj platformę nogami, unikając przeprostu w kolanach. UWAGA: Nie odrywaj lędźwi od oparcia maszyny.", "czas_min": 3, "parametry": "4x12", "miesnie": "Czworogłowe, pośladki"},
        {"nazwa": "Wyprosty nóg na maszynie", "opis": "POZYCJA: Siedząc, wałek pod stawami skokowymi. RUCH: Prostuj kolana, napinając mięśnie ud. UWAGA: Ruch wykonuj płynnie, bez gwałtownych szarpnięć ciężarem.", "czas_min": 2, "parametry": "3x15", "miesnie": "Czworogłowe"},
        {"nazwa": "Uginanie nóg na maszynie", "opis": "POZYCJA: Leżąc przodem, wałek nad piętami. RUCH: Zginaj kolana, przyciągając pięty do pośladków. UWAGA: Biodra powinny cały czas przylegać do ławki maszyny.", "czas_min": 2, "parametry": "3x15", "miesnie": "Dwugłowe uda"},
        {"nazwa": "Przysiad Goblet (z kettlem)", "opis": "POZYCJA: Kettlebell trzymany oburącz przy klatce piersiowej. RUCH: Wykonuj głęboki przysiad. UWAGA: Kettlebell pomaga utrzymać pionową sylwetkę, nie garb się.", "czas_min": 3, "parametry": "3x12", "miesnie": "Czworogłowe, core"},
        {"nazwa": "Przysiad bułgarski", "opis": "POZYCJA: Jedna noga oparta o ławkę w tyle. RUCH: Przysiad na nodze wykrocznej. UWAGA: Bardzo wymagające ćwiczenie na stabilizację, pilnuj kolana.", "czas_min": 3, "parametry": "3x10 na nogę", "miesnie": "Czworogłowe, pośladki"},
        {"nazwa": "Wspięcia na palce (maszyna)", "opis": "POZYCJA: Stopy na krawędzi platformy. RUCH: Unoszenie pięt jak najwyżej. UWAGA: Pełen zakres ruchu – rozciągnięcie w dole i spięcie w górze.", "czas_min": 2, "parametry": "4x20", "miesnie": "Łydki"},
        {"nazwa": "Kettlebell Swing", "opis": "POZYCJA: Stojąc, kettlebell między nogami. RUCH: Dynamiczny wymach bioder w przód. UWAGA: Siła ruchu płynie z bioder, a nie z barków.", "czas_min": 3, "parametry": "3x15", "miesnie": "Dwugłowe, pośladki, core"},
        {"nazwa": "Wykroki z gumą oporową", "opis": "POZYCJA: Guma pod przednią stopą, końce w dłoniach. RUCH: Klasyczny wykrok z dodatkowym oporem gumy. UWAGA: Kontroluj tempo fazy ekscentrycznej (opuszczania).", "czas_min": 3, "parametry": "3x12", "miesnie": "Czworogłowe, pośladki"},
        {"nazwa": "Przysiad Sumo", "opis": "POZYCJA: Bardzo szeroki rozkrok, palce stóp na zewnątrz. RUCH: Przysiad z ciężarem w dłoniach. UWAGA: Plecy proste, większy nacisk na wewnętrzną część ud.", "czas_min": 3, "parametry": "3x12", "miesnie": "Przywodziciele, czworogłowe"},
        {"nazwa": "Zakroki", "opis": "POZYCJA: Stojąc, wykonaj krok w tył. RUCH: Obniżaj biodra, aż kolano tylnej nogi dotknie ziemi. UWAGA: Zachowaj stabilny tułów, nie chwiej się.", "czas_min": 3, "parametry": "3x10 na nogę", "miesnie": "Czworogłowe, pośladki"},
        {"nazwa": "Wyprosty nóg w siadzie z gumą", "opis": "POZYCJA: Siedząc, guma zaczepiona o nogę krzesła i staw skokowy. RUCH: Prostuj nogę. UWAGA: Powolna faza powrotu, maksymalne napięcie uda.", "czas_min": 2, "parametry": "3x20", "miesnie": "Czworogłowe"},
        {"nazwa": "Martwy ciąg na jednej nodze", "opis": "POZYCJA: Stojąc na jednej nodze, hantel w dłoni. RUCH: Skłon z jednoczesnym wyprostowaniem drugiej nogi w tył. UWAGA: Ćwiczenie na równowagę, skup się na punkcie przed sobą.", "czas_min": 3, "parametry": "3x10 na nogę", "miesnie": "Dwugłowe, pośladki, core"},
        {"nazwa": "Wchodzenie na skrzynię (Step-up)", "opis": "POZYCJA: Stojąc przed skrzynią. RUCH: Wejdź na skrzynię, prostując nogę, potem wróć. UWAGA: Cała stopa musi znajdować się na skrzyni.", "czas_min": 3, "parametry": "3x12 na nogę", "miesnie": "Pośladki, czworogłowe"},
        {"nazwa": "Sissy Squat", "opis": "POZYCJA: Stojąc, trzymając się stabilnego punktu. RUCH: Odchyl tułów w tył, wypychając kolana daleko w przód. UWAGA: Tylko dla osób bez problemów ze stawami kolanowymi.", "czas_min": 2, "parametry": "3x10", "miesnie": "Czworogłowe"},
        {"nazwa": "Wykroki chodzone z hantlami", "opis": "POZYCJA: Stojąc z hantlami. RUCH: Wykonuj naprzemienne wykroki idąc przed siebie. UWAGA: Pilnuj stabilności, nie pozwól kolanom uciekać do wewnątrz.", "czas_min": 3, "parametry": "3x20m", "miesnie": "Czworogłowe, pośladki"},
        {"nazwa": "Przysiad z gumą nad kolanami", "opis": "POZYCJA: Guma nad kolanami. RUCH: Wykonuj przysiad, cały czas rozpychając gumę kolanami. UWAGA: To ćwiczenie uczy prawidłowej pracy kolan.", "czas_min": 3, "parametry": "3x15", "miesnie": "Nogi, pośladki średnie"},
        {"nazwa": "Wspięcia na palce bez obciążenia", "opis": "POZYCJA: Stojąc, ręce na biodrach. RUCH: Dynamiczne wspięcia na palce. UWAGA: Duża objętość dla poprawy wytrzymałości mięśni łydki.", "czas_min": 2, "parametry": "3x30", "miesnie": "Łydki"}
    ],
   "Pośladki": [
        {"nazwa": "Hip Thrust (ze sztangą)", "opis": "POZYCJA: Łopatki oparte o stabilną ławkę, sztanga na biodrach. RUCH: Wypchnij biodra w górę, spinając pośladki w szczycie ruchu. UWAGA: Nie wyginaj odcinka lędźwiowego – ruch zachodzi w biodrach, nie w plecach.", "czas_min": 4, "parametry": "4x10", "miesnie": "Pośladkowy wielki"},
        {"nazwa": "Odwodzenie nogi na wyciągu", "opis": "POZYCJA: Stojąc, opaska na kostce, trzymaj się wyciągu. RUCH: Odprowadzaj prostą nogę w tył, kontrolując napięcie pośladka. UWAGA: Tułów stabilny, nie wykonuj zamachu tułowiem.", "czas_min": 2, "parametry": "3x15", "miesnie": "Pośladkowy wielki"},
        {"nazwa": "Glute Bridge z gumą", "opis": "POZYCJA: Leżenie tyłem, guma nad kolanami. RUCH: Unoś biodra, rozpychając kolana na zewnątrz przeciwko oporowi gumy. UWAGA: Stopy mocno wciśnięte w ziemię.", "czas_min": 2, "parametry": "3x20", "miesnie": "Pośladkowy średni, wielki"},
        {"nazwa": "Clamshell (Muszelka)", "opis": "POZYCJA: Leżenie na boku, kolana ugięte 90st. RUCH: Unoś górne kolano, nie rozłączając stóp. UWAGA: Miednica musi być nieruchoma, nie odchylaj się w tył.", "czas_min": 2, "parametry": "3x20", "miesnie": "Pośladkowy średni"},
        {"nazwa": "Martwy ciąg Rumuński", "opis": "POZYCJA: Sztanga/hantle w dłoniach, lekko ugięte kolana. RUCH: Wypychaj biodra mocno w tył, aż poczujesz rozciąganie pośladków. UWAGA: Plecy proste, ciężar blisko nóg.", "czas_min": 3, "parametry": "3x12", "miesnie": "Pośladki, dwugłowe"},
        {"nazwa": "Kabel kickbacks (wypychanie)", "opis": "POZYCJA: W klęku podpartym, opaska wyciągu na stopie. RUCH: Wypychaj piętę w górę w stronę sufitu. UWAGA: Nie wyginaj nadmiernie kręgosłupa w lędźwiach.", "czas_min": 3, "parametry": "3x15", "miesnie": "Pośladkowy wielki"},
        {"nazwa": "Wspięcia na skrzynię (Step-up)", "opis": "POZYCJA: Stań przed skrzynią. RUCH: Wejdź na skrzynię, mocno napinając pośladek nogi, która wykonuje pracę. UWAGA: Nie odpychaj się nogą znajdującą się na ziemi.", "czas_min": 3, "parametry": "3x10 na nogę", "miesnie": "Pośladki, nogi"},
        {"nazwa": "Fire Hydrant z gumą", "opis": "POZYCJA: W klęku podpartym, guma nad kolanami. RUCH: Odwodź nogę ugiętą w kolanie do boku (jak pies przy hydrancie). UWAGA: Ruch w biodrze, tułów nieruchomo.", "czas_min": 2, "parametry": "3x15", "miesnie": "Pośladkowy średni"},
        {"nazwa": "Przysiad Sumo (z kettlem)", "opis": "POZYCJA: Bardzo szeroki rozkrok, palce stóp na zewnątrz. RUCH: Przysiad z kettlem między nogami. UWAGA: Pilnuj, aby kolana nie schodziły do środka.", "czas_min": 3, "parametry": "3x12", "miesnie": "Pośladki, przywodziciele"},
        {"nazwa": "Abdukcja na maszynie", "opis": "POZYCJA: Siedząc na maszynie, rozpychaj poduszki na zewnątrz. RUCH: Powolny ruch rozwodzenia nóg. UWAGA: W fazie powrotu kontroluj ciężar, nie pozwól mu uderzyć o stos.", "czas_min": 2, "parametry": "3x15", "miesnie": "Pośladkowy średni"},
        {"nazwa": "Hip Thrust na 1 nodze", "opis": "POZYCJA: Jedna noga w górze, druga na podłożu. RUCH: Wypychanie bioder. UWAGA: To wersja zaawansowana, wymaga idealnej stabilności miednicy.", "czas_min": 3, "parametry": "3x10", "miesnie": "Pośladki"},
        {"nazwa": "Wykroki boczne (Side Lunge)", "opis": "POZYCJA: Stojąc, szeroki krok w bok. RUCH: Przenieś ciężar na nogę wykroczną, wypychając biodro w tył. UWAGA: Noga, która zostaje w miejscu, jest wyprostowana.", "czas_min": 3, "parametry": "3x12", "miesnie": "Pośladki, uda"},
        {"nazwa": "Mostki biodrowe na 1 nodze", "opis": "POZYCJA: Leżenie tyłem, jedna noga uniesiona. RUCH: Unoś biodra, używając tylko jednej nogi. UWAGA: Utrzymuj miednicę w poziomie przez cały ruch.", "czas_min": 2, "parametry": "3x12", "miesnie": "Pośladki"},
        {"nazwa": "Kettlebell Swing", "opis": "POZYCJA: Stojąc, kettlebell między nogami. RUCH: Wykonaj dynamiczny wymach biodrami. UWAGA: Plecy proste, ruch inicjują biodra, nie ręce.", "czas_min": 3, "parametry": "3x20", "miesnie": "Pośladki, dwugłowe"},
        {"nazwa": "Odwodzenie w leżeniu na boku", "opis": "POZYCJA: Leżenie na boku, noga prosta. RUCH: Unoś prostą nogę w górę. UWAGA: Palce stopy skierowane do przodu, nie do sufitu.", "czas_min": 2, "parametry": "3x20", "miesnie": "Pośladkowy średni"},
        {"nazwa": "Good Morning", "opis": "POZYCJA: Sztanga na karku, stopy na szerokość bioder. RUCH: Skłon w przód z wypchnięciem bioder w tył. UWAGA: Tylko do poziomu równoległego do podłogi.", "czas_min": 3, "parametry": "3x10", "miesnie": "Pośladki, prostowniki"},
        {"nazwa": "Przysiad w tempie (pauza)", "opis": "POZYCJA: Klasyczna. RUCH: Zejdź w dół (3 sekundy), zatrzymaj na 2 sekundy, wróć szybko. UWAGA: Stałe napięcie mięśniowe w pauzie.", "czas_min": 3, "parametry": "3x10", "miesnie": "Pośladki"},
        {"nazwa": "Glute Machine (maszyna)", "opis": "POZYCJA: Stojąc lub w klęku na maszynie. RUCH: Wypychanie nogi w tył przeciwko oporowi. UWAGA: Nie pogłębiaj lędźwi podczas ruchu.", "czas_min": 3, "parametry": "3x12", "miesnie": "Pośladkowy wielki"},
        {"nazwa": "Spacer kraba z gumą", "opis": "POZYCJA: Pozycja półprzysiadu, guma nad kolanami. RUCH: Krok w bok, utrzymując napięcie taśmy. UWAGA: Nie prostuj nóg w trakcie chodzenia.", "czas_min": 2, "parametry": "3x20m", "miesnie": "Pośladkowy średni"},
        {"nazwa": "Unoszenie nogi w klęku (z gumą)", "opis": "POZYCJA: Klęk, guma pod stopą i dłonią. RUCH: Unoś nogę do boku/tyłu. UWAGA: Stabilny tułów.", "czas_min": 2, "parametry": "3x15", "miesnie": "Pośladki"}
    ],
    "Zakończenie treningu": [
        {"nazwa": "Schładzanie na rowerku stacjonarnym", "opis": "POZYCJA: Siedząc na rowerku. RUCH: Spokojne pedałowanie z bardzo niskim oporem. UWAGA: Skup się na głębokim, wolnym oddechu, aby uspokoić tętno.", "czas_min": 5, "parametry": "5 minut", "miesnie": "Całe ciało, układ krążenia"},
        {"nazwa": "Rozciąganie taśmy tylnej w siadzie", "opis": "POZYCJA: Siad prosty, nogi złączone. RUCH: Wykonaj skłon w przód, próbując chwycić stopy. UWAGA: Nie szarp, pogłębiaj zakres z każdym wydechem.", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Grupa kulszowo-goleniowa, łydki"},
        {"nazwa": "Pozycja dziecka (Child's Pose)", "opis": "POZYCJA: Klęk na podłodze, pośladki na piętach, ręce wyciągnięte w przód. RUCH: Opuść czoło do maty. UWAGA: Rozluźnij mięśnie grzbietu i poczuj wydłużanie kręgosłupa.", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Prostowniki grzbietu, barki"},
        {"nazwa": "Rozciąganie klatki piersiowej przy ścianie", "opis": "POZYCJA: Stań bokiem do ściany, oprzyj przedramię na niej. RUCH: Skręć tułów w przeciwną stronę. UWAGA: Poczuj rozciąganie w obrębie klatki i przedniej części barku.", "czas_min": 2, "parametry": "1 minuta na stronę", "miesnie": "Klatka piersiowa, barki"},
        {"nazwa": "Pozycja Sfinksa", "opis": "POZYCJA: Leżenie przodem, podparcie na przedramionach. RUCH: Wypychaj klatkę w przód i w górę. UWAGA: Jeśli czujesz ból w lędźwiach, obniż tułów – to nie jest ćwiczenie siłowe.", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Brzuch, prostowniki grzbietu"},
        {"nazwa": "Rozciąganie zginaczy bioder", "opis": "POZYCJA: Wykrok, tylne kolano na macie. RUCH: Wypchnij biodra w przód. UWAGA: Napnij pośladek nogi zakrocznej, aby zwiększyć rozciąganie z przodu uda.", "czas_min": 2, "parametry": "1 minuta na stronę", "miesnie": "Zginacze bioder, czworogłowe"},
        {"nazwa": "Kot-krowa (rozluźnienie)", "opis": "POZYCJA: Klęk podparty. RUCH: Naprzemienne wyginanie kręgosłupa w górę i w dół. UWAGA: Ruch w tempie oddechu – wydech w wygięciu, wdech w obniżeniu.", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Cały kręgosłup"},
        {"nazwa": "Rozciąganie pośladków (Gołąb)", "opis": "POZYCJA: Siad ze zgiętą nogą przed sobą. RUCH: Pochyl tułów nad nogą. UWAGA: Utrzymuj biodra w linii, nie przechylaj się na boki.", "czas_min": 2, "parametry": "1 minuta na stronę", "miesnie": "Pośladki, rotator biodra"},
        {"nazwa": "Skręt tułowia w leżeniu", "opis": "POZYCJA: Leżenie tyłem, ręce szeroko. RUCH: Przenieś zgięte kolana w bok, dotykając nimi podłogi. UWAGA: Barki muszą przylegać do maty.", "czas_min": 2, "parametry": "1 minuta na stronę", "miesnie": "Mięśnie skośne brzucha, plecy"},
        {"nazwa": "Rozciąganie łydki przy ścianie", "opis": "POZYCJA: Stanie przy ścianie, stopa w tyle. RUCH: Dociśnij piętę do ziemi. UWAGA: Kolano nogi zakrocznej powinno być wyprostowane.", "czas_min": 2, "parametry": "1 minuta na stronę", "miesnie": "Brzuchaty łydki, płaszczkowaty"},
        {"nazwa": "Pozycja 'Pies z głową w dół'", "opis": "POZYCJA: Podpór, wypchnij biodra wysoko w górę. RUCH: Staraj się dociągnąć pięty do ziemi. UWAGA: Plecy muszą być proste.", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Tył nóg, plecy"},
        {"nazwa": "Rozciąganie bicepsa przy ścianie", "opis": "POZYCJA: Dłoń oparta o ścianę, kciuk w górę. RUCH: Obróć tułów w przeciwną stronę. UWAGA: Poczuj rozciąganie wzdłuż ramienia.", "czas_min": 2, "parametry": "1 minuta na stronę", "miesnie": "Biceps, przedramiona"},
        {"nazwa": "Rozciąganie tricepsa za głową", "opis": "POZYCJA: Stojąc, zegnij rękę w łokciu. RUCH: Drugą ręką dociśnij łokieć do dołu. UWAGA: Nie wyginaj kręgosłupa, trzymaj brzuch napięty.", "czas_min": 2, "parametry": "1 minuta na stronę", "miesnie": "Triceps"},
        {"nazwa": "Wydłużanie szyi (boczne)", "opis": "POZYCJA: Siedząc prosto. RUCH: Przyciągnij ucho do barku, przeciwną rękę wyciągnij w dół. UWAGA: Nie szarp głową.", "czas_min": 2, "parametry": "1 minuta na stronę", "miesnie": "Mięśnie szyi, trapezius"},
        {"nazwa": "Pozycja 'Motyla' (stopy razem)", "opis": "POZYCJA: Siad, podeszwy stóp stykają się. RUCH: Delikatnie dociskaj kolana do podłogi. UWAGA: Plecy wyprostowane.", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Przywodziciele ud"},
        {"nazwa": "Rozciąganie powięzi czworogłowego", "opis": "POZYCJA: Klęk na jednej nodze, pięta do pośladka. RUCH: Przyciągnij piętę ręką. UWAGA: Napnij pośladek, aby chronić lędźwie.", "czas_min": 2, "parametry": "1 minuta na stronę", "miesnie": "Czworogłowe"},
        {"nazwa": "Pozycja trupa (Savasana)", "opis": "POZYCJA: Leżenie tyłem, ręce wzdłuż ciała. RUCH: Pełne rozluźnienie. UWAGA: Obserwuj oddech przez 3 minuty, nic więcej nie rób.", "czas_min": 3, "parametry": "3 minuty", "miesnie": "Całe ciało (relaksacja)"},
        {"nazwa": "Krążenia nadgarstków i stawów skokowych", "opis": "POZYCJA: Siad lub stanie. RUCH: Powolne, obszerne krążenia stawami. UWAGA: Jeśli czujesz 'strzelanie', wykonuj ruch jeszcze wolniej.", "czas_min": 2, "parametry": "2 minuty", "miesnie": "Stawy obwodowe"},
        {"nazwa": "Rozciąganie najszerszego grzbietu (chwyt za słup)", "opis": "POZYCJA: Chwyć słup/framugę drzwi, odchyl biodra w tył. RUCH: Poczuj rozciąganie w boku pleców. UWAGA: Nie ciągnij na siłę.", "czas_min": 2, "parametry": "1 minuta na stronę", "miesnie": "Najszerszy grzbietu"},
        {"nazwa": "Głębokie oddychanie przeponowe", "opis": "POZYCJA: Leżenie, dłonie na brzuchu. RUCH: Wdech przez nos (brzuch rośnie), wydech przez usta. UWAGA: Wydech powinien być dwa razy dłuższy niż wdech.", "czas_min": 3, "parametry": "3 minuty", "miesnie": "Przepona"}
    ]
}

PLIK_WLASNYCH_CWICZEN = "wlasne_cwiczenia.json"

def zaladuj_wlasne_cwiczenia():
    if os.path.exists(PLIK_WLASNYCH_CWICZEN):
        try:
            with open(PLIK_WLASNYCH_CWICZEN, "r") as f:
                dane = json.load(f)
                for kat, lista in dane.get("FIZJO", {}).items():
                    if kat in BAZA_FIZJO: BAZA_FIZJO[kat].extend(lista)
                for kat, lista in dane.get("GYM", {}).items():
                    if kat in BAZA_SILOWNIA: BAZA_SILOWNIA[kat].extend(lista)
        except: pass

zaladuj_wlasne_cwiczenia()
GLOBALNA_BAZA = {**BAZA_FIZJO, **BAZA_SILOWNIA}

# AUTOMATYCZNE SORTOWANIE ALFABETYCZNE BAZ
for compliance_baza in [BAZA_FIZJO, BAZA_SILOWNIA, GLOBALNA_BAZA]:
    for compliance_kat in compliance_baza:
        compliance_baza[compliance_kat] = sorted(compliance_baza[compliance_kat], key=lambda compliance_x: compliance_x['nazwa'].lower())
# --- Wczytywanie bazy protokołów klinicznych ---
@st.cache_data
def wczytaj_protokoly():
    if os.path.exists("protokoly.json"):
        try:
            with open("protokoly.json", "r", encoding="utf-8") as f: return json.load(f)
        except: pass
    return {}

PROTOKOLY_KLINICZNE = wczytaj_protokoly()

def generuj_protokol(nazwa_choroby):
    plan = []
    naglowek = {"nazwa": f"PROTOKÓŁ: {nazwa_choroby}", "typ": "Kliniczny", "partie": "-", "opis": "Zestaw ułożony celowo pod jednostkę chorobową.", "czas_min": 0, "parametry": "-", "miesnie": "-", "uwagi": ""}
    plan.append(("NAGŁÓWEK DNIA", naglowek))
    
    lista_cwiczen_w_chorobie = []
    for kategoria, choroby in PROTOKOLY_KLINICZNE.items():
        if nazwa_choroby in choroby:
            lista_cwiczen_w_chorobie = choroby[nazwa_choroby]
            break
    
    for nazwa_cw in lista_cwiczen_w_chorobie:
        znaleziono = False
        for kat, lista in GLOBALNA_BAZA.items():
            for cw in lista:
                if cw["nazwa"] == nazwa_cw:
                    cw_kopia = cw.copy()
                    cw_kopia["uwagi"] = ""
                    etykieta = f"GYM: {kat}" if kat in BAZA_SILOWNIA and kat not in BAZA_FIZJO else kat
                    plan.append((etykieta, cw_kopia))
                    znaleziono = True
                    break
            if znaleziono: break
            
    st.session_state.wylosowany_plan_cache = plan
    st.session_state.is_gym = False

# ==============================================================================
# LOGIKA GENERATORÓW
# ==============================================================================
def generuj_plan(profil, budzet, dni):
    plan = []
    b_fizjo = {k: list(v) for k, v in BAZA_FIZJO.items()}
    b_gym = {k: list(v) for k, v in BAZA_SILOWNIA.items()}

    def pop_random(baza, kat):
        # NOWOŚĆ: Jeśli wyczerpaliśmy wszystkie ćwiczenia dla tej partii, odnawiamy pulę z bazy głównej (tasujemy talię)
        if kat in baza and len(baza[kat]) == 0:
            if kat in BAZA_FIZJO:
                baza[kat] = list(BAZA_FIZJO[kat])
            elif kat in BAZA_SILOWNIA:
                baza[kat] = list(BAZA_SILOWNIA[kat])

        if baza.get(kat) and len(baza[kat]) > 0:
            idx = random.randint(0, len(baza[kat])-1)
            cw = baza[kat].pop(idx).copy()
            cw["uwagi"] = ""
            return cw
        return None

    czy_split = "Split" in profil
    czy_wielo_dniowy = czy_split or "Kompleksowy" in profil
    rzeczywista_liczba_dni = dni if czy_wielo_dniowy else 1
    is_gym = "GYM:" in profil

    dni_tygodnia = ["Poniedziałek", "Wtorek", "Środa", "Czwartek", "Piątek", "Sobota", "Niedziela"]
    
    if czy_split:
        # AKTUALIZACJA: Dodano Barki i Brzuch / Core do cyklu
        kolejnosc_gym = ["Klatka piersiowa", "Plecy", "Nogi", "Barki", "Ręce", "Pośladki", "Brzuch / Core"]
        kolejnosc_fizjo = ["Głowa/Szyja", "Kończyna górna", "Core (Tułów)", "Kończyna dolna"]
        
        # reszta kodu pozostaje bez zmian...

        plan_na_dni = []
        for i in range(rzeczywista_liczba_dni):
            if is_gym:
                # Wyciągamy jedną partię, zapętlając listę po jej zakończeniu
                partia = kolejnosc_gym[i % len(kolejnosc_gym)]
            else:
                partia = kolejnosc_fizjo[i % len(kolejnosc_fizjo)]
            
            # Wrzucamy do planu dokładnie jedną partię dla danego dnia
            plan_na_dni.append([partia])

    for i in range(rzeczywista_liczba_dni):
        if czy_wielo_dniowy:
            if czy_split:
                dzien = dni_tygodnia[i % 7]
                partie = plan_na_dni[i]
                partie_str = ' + '.join([p.upper() for p in partie])
                nazwa_dnia = f"{dzien}: {partie_str}"
                naglowek = {"nazwa": nazwa_dnia, "typ": "Split", "partie": partie_str, "opis": "", "czas_min": 0, "parametry": "-", "miesnie": "-", "uwagi": ""}
            else:
                nazwa_dnia = f"Dzień {i+1} (FIZJO: Kompleksowy)"
                naglowek = {"nazwa": nazwa_dnia, "typ": "Kompleksowy", "partie": "", "opis": "", "czas_min": 0, "parametry": "-", "miesnie": "-", "uwagi": ""}
            plan.append(("NAGŁÓWEK DNIA", naglowek))

        if is_gym:
            cw_start = pop_random(b_gym, "Rozgrzewka") or random.choice(BAZA_SILOWNIA["Rozgrzewka"]).copy()
            cw_start["uwagi"] = ""
            plan.append(("GYM: Rozgrzewka", cw_start))
        else:
            cw_start = pop_random(b_fizjo, "Oddechowe") or random.choice(BAZA_FIZJO["Oddechowe"]).copy()
            cw_start["uwagi"] = ""
            plan.append(("Oddechowe (Rozgrzewka)", cw_start))
        
        realny_czas = cw_start.get("czas_min", 2) if not is_gym else 0

        if czy_split:
            partie = plan_na_dni[i]
            
            if is_gym:
                # Moduł GYM: Budżet dzielimy równo na liczbę partii w danym dniu.
                # (Jeśli masz budżet 6 ćwiczeń i 2 partie w dniu, wylosuje po 3 na każdą)
                budzet_na_partie = max(1, budzet // len(partie))
                for p in partie:
                    dodano = 0
                    while dodano < budzet_na_partie:
                        cw = pop_random(b_gym, p)
                        if not cw: break
                        plan.append((f"GYM: {p}", cw))
                        dodano += 1
            else:
                # Moduł FIZJO: Budżet to czas w minutach! Dzielimy wolny czas na liczbę partii.
                czas_koncowy = 2
                budzet_czasu_na_partie = max(1, (budzet - realny_czas) // len(partie))
                for p in partie:
                    wykorzystany_czas_partii = 0
                    while wykorzystany_czas_partii + czas_koncowy < budzet_czasu_na_partie:
                        dostepne = [c for c in b_fizjo.get(p, []) if wykorzystany_czas_partii + c["czas_min"] + czas_koncowy <= budzet_czasu_na_partie]
                        if not dostepne: break
                        cw_wybrane = dostepne[random.randint(0, len(dostepne)-1)]
                        cw = b_fizjo[p].pop(b_fizjo[p].index(cw_wybrane)).copy()
                        cw["uwagi"] = ""
                        plan.append((p, cw))
                        wykorzystany_czas_partii += cw["czas_min"]
                        realny_czas += cw["czas_min"]
        else:
            if is_gym:
                if "Ogólnorozwojowy" in profil:
                    # AKTUALIZACJA: Pełen przegląd partii FBW
                    for p in ["Klatka piersiowa", "Plecy", "Nogi", "Barki", "Ręce", "Pośladki", "Brzuch / Core"]:
                        dodano = 0
                        while dodano < budzet:
                            cw = pop_random(b_gym, p)
                            if not cw: break
                            plan.append((f"GYM: {p}", cw))
                            dodano += 1
                else:
                    kat = profil.split(" - ")[1]
                    dodano = 0
                    while dodano < budzet:
                        cw = pop_random(b_gym, kat)
                        if not cw: break
                        plan.append((f"GYM: {kat}", cw))
                        dodano += 1
            else:
                czas_koncowy = 2 
                if "Kompleksowy" in profil:
                    lancuch = ["Głowa/Szyja", "Kończyna górna", "Core (Tułów)", "Kończyna dolna"]
                    puste = 0
                    while realny_czas + czas_koncowy < budzet and puste < len(lancuch):
                        dodano_cos = False
                        for kat in lancuch:
                            dostepne = [c for c in b_fizjo.get(kat, []) if realny_czas + c["czas_min"] + czas_koncowy <= budzet]
                            if dostepne:
                                cw_wybrane = dostepne[random.randint(0, len(dostepne)-1)]
                                cw = b_fizjo[kat].pop(b_fizjo[kat].index(cw_wybrane)).copy()
                                cw["uwagi"] = ""
                                plan.append((kat, cw))
                                realny_czas += cw["czas_min"]
                                dodano_cos = True
                        if not dodano_cos: puste += 1
                else:
                    kat = profil.split(" - ")[1].replace("Tylko ", "")
                    while realny_czas + czas_koncowy < budzet:
                        dostepne = [c for c in b_fizjo.get(kat, []) if realny_czas + c["czas_min"] + czas_koncowy <= budzet]
                        if not dostepne: break
                        cw_wybrane = dostepne[random.randint(0, len(dostepne)-1)]
                        cw = b_fizjo[kat].pop(b_fizjo[kat].index(cw_wybrane)).copy()
                        cw["uwagi"] = ""
                        plan.append((kat, cw))
                        realny_czas += cw["czas_min"]

        if is_gym:
            cw_koniec = pop_random(b_gym, "Zakończenie treningu") or random.choice(BAZA_SILOWNIA["Zakończenie treningu"]).copy()
            cw_koniec["uwagi"] = ""
            plan.append(("GYM: Zakończenie", cw_koniec))
        else:
            cw_koniec = pop_random(b_fizjo, "Oddechowe") or random.choice(BAZA_FIZJO["Oddechowe"]).copy()
            cw_koniec["uwagi"] = ""
            plan.append(("Oddechowe (Wyciszenie)", cw_koniec))

    st.session_state.wylosowany_plan_cache = plan
    st.session_state.is_gym = is_gym

# ==============================================================================
# REJESTR EKSPORTÓW
# ==============================================================================
def generuj_docx():
    doc = Document()
    doc.add_heading('ZINTEGROWANA KARTA REHABILITACYJNO-TRENINGOWA', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, (kat, cw) in enumerate(st.session_state.wylosowany_plan_cache, 1):
        if kat == "NAGŁÓWEK DNIA":
            p = doc.add_paragraph()
            run = p.add_run(f"\n{cw['nazwa']}\n")
            run.bold = True; run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        p = doc.add_paragraph()
        p.add_run(f"{idx}. {kat.upper()}: {cw['nazwa']}\n").bold = True
        p.add_run("DAWKOWANIE: ").bold = True
        p.add_run(f"{cw['parametry']}\n")
        p.add_run("MIESNIE: ").bold = True
        p.add_run(f"{cw['miesnie']}\n")
        p.add_run("OPIS: ").bold = True
        p.add_run(f"{cw['opis']}\n")
        p.add_run("UWAGI: ").bold = True
        p.add_run(f"{cw.get('uwagi', '..........................................................')}\n")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def generuj_excel_fizjo(liczba_dni):
    if not st.session_state.wylosowany_plan_cache: 
        return None
        
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Plan_Standardowy"
    
    thin_border = Border(left=Side(style='thin', color='000000'), right=Side(style='thin', color='000000'), 
                         top=Side(style='thin', color='000000'), bottom=Side(style='thin', color='000000'))
    
    # Pancerny format ARGB (wymuszany przez MS Excel i Google Sheets)
    fill_day = PatternFill(start_color="FF28A745", end_color="FF28A745", fill_type="solid")
    fill_header = PatternFill(start_color="FFD4EDDA", end_color="FFD4EDDA", fill_type="solid")
    font_white = Font(color="FFFFFFFF", bold=True, size=11)
    
    # Nagłówek główny dokumentu
    ws.cell(row=1, column=1, value="KARTA PACJENTA / PLAN TRENINGOWY").font = Font(bold=True, size=14, color="FF28A745")
    ws.cell(row=2, column=1, value="Imię i Nazwisko:").font = Font(bold=True)
    ws.cell(row=2, column=4, value="Data:").font = Font(bold=True)
    ws.cell(row=3, column=1, value="Status / Uwagi główne:").font = Font(bold=True)
    
    realny_czas = 0
    for k, c in st.session_state.wylosowany_plan_cache:
        if k != "NAGŁÓWEK DNIA" and isinstance(c, dict):
            realny_czas += c.get('czas_min', 0)
            
    ws.cell(row=5, column=1, value="Całkowity szacowany czas terapii/treningu:").font = Font(bold=True)
    ws.cell(row=5, column=2, value=f"{realny_czas} min")
    
    # --- NOWOŚĆ: Dynamiczne wykrywanie struktury planu ---
    ma_naglowki_dni = any(kat == "NAGŁÓWEK DNIA" for kat, _ in st.session_state.wylosowany_plan_cache)
    
    row_idx = 7
    dzien_aktualny = 0
    lp = 1
    ostatnia_kategoria = None
    pierwszy_wiersz_bez_naglowka = not ma_naglowki_dni
    
    for kat, cw in st.session_state.wylosowany_plan_cache:
        is_new_section = False
        sekcja_tytul = ""
        
        # Warunek 1: Klasyczny plan posiadający sztywne nagłówki dni
        if kat == "NAGŁÓWEK DNIA":
            is_new_section = True
            dzien_aktualny += 1
            sekcja_tytul = cw.get('nazwa', f'Dzień {dzien_aktualny}') if isinstance(cw, dict) else str(cw)
        
        # Warunek 2: Plan ukierunkowany - generowanie pierwszego nagłówka na starcie danych
        elif pierwszy_wiersz_bez_naglowka:
            is_new_section = True
            sekcja_tytul = str(kat) if kat else "Zestaw Ćwiczeń"
            ostatnia_kategoria = kat
            pierwszy_wiersz_bez_naglowka = False
            
        # Warunek 3: Plan ukierunkowany - wykrycie zmiany bloku / kategorii ćwiczeń
        elif not ma_naglowki_dni and kat != ostatnia_kategoria:
            is_new_section = True
            sekcja_tytul = str(kat) if kat else "Kolejna Sekcja"
            ostatnia_kategoria = kat
            
        # Rysowanie struktury nagłówkowej, jeśli spełniony został jeden z powyższych warunków
        if is_new_section:
            if row_idx > 7: 
                row_idx += 1  # Estetyczny dodatkowy odstęp między tabelami
            
            # Zielony pasek tytułowy sekcji
            ws.merge_cells(start_row=row_idx, start_column=1, end_row=row_idx, end_column=6)
            c_day = ws.cell(row=row_idx, column=1, value=f"📋 {sekcja_tytul}")
            c_day.font = font_white
            c_day.fill = fill_day
            c_day.alignment = Alignment(vertical='center', horizontal='left')
            for i in range(1, 7): 
                ws.cell(row=row_idx, column=i).border = thin_border
            row_idx += 1
            
            # Wiersz z brakującymi dotychczas nazwami kolumn (Jasnozielony)
            naglowki = ["L.p.", "Procedura / Ćwiczenie", "Czas", "Dawkowanie / Parametry", "Anatomia / Cel", "Instrukcja wykonania"]
            for i, nagl in enumerate(naglowki, 1):
                c_h = ws.cell(row=row_idx, column=i, value=nagl)
                c_h.font = Font(bold=True)
                c_h.fill = fill_header
                c_h.border = thin_border
                c_h.alignment = Alignment(horizontal="center", vertical="center")
            row_idx += 1
            lp = 1  # Resetowanie licznika pozycji wewnątrz nowo otwartej sekcji
            
        if kat == "NAGŁÓWEK DNIA":
            continue
            
        # Wiersze z danymi ćwiczenia
        if isinstance(cw, dict):
            n_cw = cw.get('nazwa', 'Brak')
            c_min = f"{cw.get('czas_min', 0)} min"
            param = cw.get('parametry', '-')
            mies = cw.get('miesnie', '-')
            opis = cw.get('opis', '-')
        else:
            n_cw = str(cw); c_min = "-"; param = "-"; mies = "-"; opis = "-"
            
        dane_wiersza = [lp, n_cw, c_min, param, mies, opis]
        for i, val in enumerate(dane_wiersza, 1):
            c_d = ws.cell(row=row_idx, column=i, value=val)
            c_d.border = thin_border
            c_d.alignment = Alignment(wrap_text=True, vertical='top', horizontal='left' if i > 1 else 'center')
        row_idx += 1; lp += 1

    # Automatyczne szerokości kolumn dostosowane do układu standardowego
    for c_letter, c_width in [('A', 6), ('B', 30), ('C', 10), ('D', 20), ('E', 25), ('F', 50)]:
        ws.column_dimensions[c_letter].width = c_width

    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()

def generuj_excel_gym(liczba_dni):
    if not st.session_state.wylosowany_plan_cache: return None
    plan = st.session_state.wylosowany_plan_cache

    wb = openpyxl.Workbook()
    # Definicja krawędzi
    thin_side = Side(style='thin', color='000000')
    medium_side = Side(style='medium', color='000000')
    thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
    
    # 8-znakowe pancerne kolory ARGB (z prefiksem FF)
    fill_blue_header = PatternFill(start_color="FF1F497D", end_color="FF1F497D", fill_type="solid")
    fill_day_header = PatternFill(start_color="FF007BFF", end_color="FF007BFF", fill_type="solid")
    fill_week_title = PatternFill(start_color="FFCCE5FF", end_color="FFCCE5FF", fill_type="solid")
    fill_series_title = PatternFill(start_color="FFE6F2FF", end_color="FFE6F2FF", fill_type="solid")
    font_white = Font(color="FFFFFFFF", bold=True)
    
    # --- ZAKŁADKA 1: PLAN TRENINGOWY ---
    ws_plan = wb.active
    ws_plan.title = "PLAN TRENINGOWY"
    
    row_idx = 1; dzien_idx = 1; ex_nr = 1 
    
    for kat, cw in plan:
        is_header = (kat == "NAGŁÓWEK DNIA")
        if is_header:
            if row_idx > 1: row_idx += 1
            nazwa_dnia = cw.get('nazwa', f'Dzień {dzien_idx}') if isinstance(cw, dict) else str(cw)
            ws_plan.cell(row=row_idx, column=1, value=f"Trening {dzien_idx} - {nazwa_dnia}").font = Font(bold=True, size=14, color="FF007BFF")
            row_idx += 1
            
            # Nagłówki kolumn (9 kolumn)
            naglowki_gym = ["NR", "ĆWICZENIE", "SERIE", "POWTÓRZENIA", "TEMPO", "RIR", "PRZERWA", "Wideo", "OPIS WYKONANIA"]
            for col_idx, h in enumerate(naglowki_gym, 1):
                c = ws_plan.cell(row=row_idx, column=col_idx, value=h)
                c.fill = fill_blue_header; c.font = font_white; c.border = thin_border; c.alignment = Alignment(horizontal="center", vertical="center")
            row_idx += 1; ex_nr = 1; dzien_idx += 1
        else:
            if isinstance(cw, dict):
                nazwa_cw = cw.get('nazwa', str(cw))
                parametry_str = str(cw.get('parametry', '1x1'))
                tempo_cw = cw.get('tempo', '-')
                rir_cw = cw.get('rir', '-')
                przerwa_cw = cw.get('przerwa', '-')
                wideo_cw = cw.get('wideo', '-')
                opis_cw = cw.get('opis', '-')
            else:
                nazwa_cw = str(cw); parametry_str = "1x1"; tempo_cw = "-"; rir_cw = "-"; przerwa_cw = "-"; wideo_cw = "-"; opis_cw = "-"
                
            if 'x' in parametry_str.lower():
                czlon = parametry_str.replace('X', 'x').split('x', 1)
                serie, powt = czlon[0].strip(), czlon[1].strip()
            else:
                serie, powt = "1", parametry_str.strip()
                
            # Wpisywanie danych krok po kroku do wszystkich 9 kolumn
            ws_plan.cell(row=row_idx, column=1, value=ex_nr).alignment = Alignment(horizontal="center", vertical="center")
            
            c_nazwa = ws_plan.cell(row=row_idx, column=2, value=nazwa_cw)
            c_nazwa.alignment = Alignment(wrap_text=True, vertical='center', horizontal='left')
            
            ws_plan.cell(row=row_idx, column=3, value=serie).alignment = Alignment(horizontal="center", vertical="center")
            ws_plan.cell(row=row_idx, column=4, value=powt).alignment = Alignment(horizontal="center", vertical="center")
            ws_plan.cell(row=row_idx, column=5, value=tempo_cw).alignment = Alignment(horizontal="center", vertical="center")
            ws_plan.cell(row=row_idx, column=6, value=rir_cw).alignment = Alignment(horizontal="center", vertical="center")
            ws_plan.cell(row=row_idx, column=7, value=przerwa_cw).alignment = Alignment(horizontal="center", vertical="center")
            ws_plan.cell(row=row_idx, column=8, value=wideo_cw).alignment = Alignment(horizontal="left", vertical="center")
            
            # Wpisanie opisu wykonania do kolumny 9 (I)
            c_opis = ws_plan.cell(row=row_idx, column=9, value=opis_cw)
            c_opis.alignment = Alignment(wrap_text=True, vertical='top', horizontal='left')
            
            # Nadanie obramowania komórkom
            for col_idx in range(1, 10): 
                ws_plan.cell(row=row_idx, column=col_idx).border = thin_border
                
            row_idx += 1; ex_nr += 1
            
    # Szerokości kolumn dla pierwszej zakładki
    for c_letter, c_width in [('A', 5), ('C', 10), ('D', 15), ('E', 10), ('F', 10), ('G', 12), ('H', 15), ('I', 50)]:
        ws_plan.column_dimensions[c_letter].width = c_width
    max_len_plan_b = max((len(str(ws_plan.cell(row=r, column=2).value or '')) for r in range(1, row_idx)), default=30)
    ws_plan.column_dimensions['B'].width = max(30, max_len_plan_b + 3)

    # --- ZAKŁADKA 2: DZIENNIK TRENINGOWY (Cały Rok - 52 tyg) ---
    ws_dz = wb.create_sheet("DZIENNIK TRENINGOWY")
    for w in range(52):
        start_col = 3 + w*18
        ws_dz.merge_cells(start_row=1, start_column=start_col, end_row=1, end_column=start_col+17)
        c = ws_dz.cell(row=1, column=start_col, value=f"Tydzień {w+1}")
        c.alignment = Alignment(horizontal="center", vertical="center"); c.fill = fill_week_title; c.font = Font(bold=True)
        for s in range(6):
            s_col = start_col + s*3
            ws_dz.merge_cells(start_row=2, start_column=s_col, end_row=2, end_column=s_col+2)
            c2 = ws_dz.cell(row=2, column=s_col, value=f"SERIA {s+1}")
            c2.alignment = Alignment(horizontal="center", vertical="center"); c2.fill = fill_series_title; c2.font = Font(bold=True)
            for idx_h, val_h in enumerate(["CIĘŻAR", "POWT.", "RIR"]):
                c3 = ws_dz.cell(row=3, column=s_col+idx_h, value=val_h); c3.alignment = Alignment(horizontal="center"); c3.font = Font(bold=True)
                ws_dz.column_dimensions[get_column_letter(s_col+idx_h)].width = 9

    ws_dz.column_dimensions['A'].width = 12; ws_dz.cell(row=3, column=1, value="KOLEJNOŚĆ").font = Font(bold=True); ws_dz.cell(row=3, column=2, value="ĆWICZENIE").font = Font(bold=True)
    
    r_dz = 4; ex_nr = 1; dzien_idx = 1; analiza_data = []
    for kat, cw in plan:
        if kat == "NAGŁÓWEK DNIA":
            if r_dz > 4: r_dz += 1
            nazwa_dnia = cw.get('nazwa', f'Dzień {dzien_idx}') if isinstance(cw, dict) else str(cw)
            ws_dz.merge_cells(start_row=r_dz, start_column=1, end_row=r_dz, end_column=2)
            c = ws_dz.cell(row=r_dz, column=1, value=f"DZIEŃ {dzien_idx} - {nazwa_dnia}")
            c.font = font_white; c.fill = fill_day_header; c.alignment = Alignment(horizontal="left", vertical="center")
            r_dz += 1; ex_nr = 1; dzien_idx += 1
        else:
            nazwa_cw = cw.get('nazwa', str(cw)) if isinstance(cw, dict) else str(cw)
            ws_dz.cell(row=r_dz, column=1, value=ex_nr).alignment = Alignment(horizontal="center", vertical="center")
            ws_dz.cell(row=r_dz, column=2, value=nazwa_cw)
            analiza_data.append((nazwa_cw, r_dz)); r_dz += 1; ex_nr += 1

    max_col_dz = 2 + 52 * 18
    ostatni_wiersz_siatki = max(100, r_dz + 15)
    for row in range(1, ostatni_wiersz_siatki + 1):
        for col in range(1, max_col_dz + 1):
            cell = ws_dz.cell(row=row, column=col)
            b_left = thin_side; b_right = thin_side; b_top = thin_side; b_bottom = thin_side
            if col == 2: b_right = medium_side
            if col == 3: b_left = medium_side
            if col > 2:
                if (col - 2) % 18 == 1: b_left = medium_side
                if (col - 2) % 18 == 0: b_right = medium_side
            if row == 3: b_bottom = medium_side
            cell.border = Border(left=b_left, right=b_right, top=b_top, bottom=b_bottom)
            if col == 2 and row >= 4: cell.alignment = Alignment(wrap_text=True, vertical='center', horizontal='left')

    max_len_dz_b = max((len(str(ws_dz.cell(row=r, column=2).value or '')) for r in range(4, r_dz)), default=30)
    ws_dz.column_dimensions['B'].width = max(30, max_len_dz_b + 3)

    # --- ZAKŁADKA 3: ANALIZA PROGRESU ---
    ws_an = wb.create_sheet("ANALIZA")
    def set_block(start_col, title, cols_titles, fill_color):
        ws_an.merge_cells(start_row=1, start_column=start_col, end_row=1, end_column=start_col+2)
        c = ws_an.cell(row=1, column=start_col, value=title)
        c.alignment = Alignment(horizontal="center"); c.font = Font(bold=True); c.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid"); c.border = thin_border
        for i, t in enumerate(cols_titles):
            c2 = ws_an.cell(row=2, column=start_col+i, value=t)
            c2.alignment = Alignment(horizontal="center"); c2.font = Font(bold=True, size=10); c2.border = thin_border; ws_an.column_dimensions[get_column_letter(start_col+i)].width = 14

    ws_an.column_dimensions['A'].width = 30; ws_an.merge_cells("A1:A2")
    c = ws_an.cell(row=1, column=1, value="ĆWICZENIE")
    c.alignment = Alignment(horizontal="center", vertical="center"); c.fill = fill_blue_header; c.font = font_white; c.border = thin_border
    
    for w in range(52): set_block(2 + w*3, f"Tydzień {w+1}", ["Max Ciężar", "Suma Powt.", "Zrealiz. Serie"], "FFE6F2FF")
    c_off = 2 + 52*3
    mw = [list(range(0,4)), list(range(4,8)), list(range(8,13)), list(range(13,17)), list(range(17,21)), list(range(21,26)), list(range(26,30)), list(range(30,34)), list(range(34,39)), list(range(39,43)), list(range(43,47)), list(range(47,52))]
    for m in range(12): set_block(c_off + m*3, f"Miesiąc {m+1}", ["Śr. Max Ciężar", "Suma Powt.", "Suma Serii"], "FFCCE5FF")
    c_off += 36; 
    for q in range(4): set_block(c_off + q*3, f"Kwartał {q+1}", ["Śr. Max Ciężar", "Suma Powt.", "Suma Serii"], "FF99CCFF")
    c_off += 12; 
    for h in range(2): set_block(c_off + h*3, f"Półrocze {h+1}", ["Śr. Max Ciężar", "Suma Powt.", "Suma Serii"], "FF66B2FF")
    c_off += 6; set_block(c_off, "Cały Rok", ["Śr. Max Ciężar", "Suma Powt.", "Suma Serii"], "FF3399FF")

    r_an = 3
    for nazwa, r_in_dz in analiza_data:
        ws_an.cell(row=r_an, column=1, value=nazwa).border = thin_border
        for w in range(52):
            w_cells = [f"'DZIENNIK TRENINGOWY'!{get_column_letter(3 + w*18 + s*3)}{r_in_dz}" for s in range(6)]
            p_cells = [f"'DZIENNIK TRENINGOWY'!{get_column_letter(4 + w*18 + s*3)}{r_in_dz}" for s in range(6)]
            ws_an.cell(row=r_an, column=2+w*3, value=f'=IF(COUNT({",".join(w_cells)})>0, MAX({",".join(w_cells)}), "")').border = thin_border
            ws_an.cell(row=r_an, column=2+w*3+1, value=f'=IF(SUM({",".join(p_cells)})>0, SUM({",".join(p_cells)}), "")').border = thin_border
            ws_an.cell(row=r_an, column=2+w*3+2, value=f'=IF(COUNT({",".join(w_cells)})>0, COUNT({",".join(w_cells)}), "")').border = thin_border
        def app_agg(cb, weeks):
            mx_c = [f"{get_column_letter(2 + w*3)}{r_an}" for w in weeks]
            pt_c = [f"{get_column_letter(2 + w*3 + 1)}{r_an}" for w in weeks]
            sr_c = [f"{get_column_letter(2 + w*3 + 2)}{r_an}" for w in weeks]
            ws_an.cell(row=r_an, column=cb, value=f'=IF(COUNT({",".join(mx_c)})>0, ROUND(AVERAGE({",".join(mx_c)}), 1), "")').border = thin_border
            ws_an.cell(row=r_an, column=cb+1, value=f'=IF(SUM({",".join(pt_c)})>0, SUM({",".join(pt_c)}), "")').border = thin_border
            ws_an.cell(row=r_an, column=cb+2, value=f'=IF(SUM({",".join(sr_c)})>0, SUM({",".join(sr_c)}), "")').border = thin_border
        c_off = 158
        for m in range(12): app_agg(c_off + m*3, mw[m])
        c_off += 36
        for q in range(4): app_agg(c_off + q*3, [w for m in range(q*3, q*3+3) for w in mw[m]])
        c_off += 12
        app_agg(c_off, [w for m in range(0, 6) for w in mw[m]]); app_agg(c_off + 3, [w for m in range(6, 12) for w in mw[m]])
        c_off += 6; app_agg(c_off, list(range(52)))
        r_an += 1

    bio = io.BytesIO(); wb.save(bio); return bio.getvalue()
            
    # --- [DZIENNIK TRENINGOWY I ANALIZA TAK SAMO JAK WCZEŚNIEJ, ALE UŻYWAJ 'plan_gym'] ---
            
    for c_letter, c_width in [('A', 5), ('C', 10), ('D', 15), ('E', 10), ('F', 10), ('G', 12), ('H', 15)]:
        ws_plan.column_dimensions[c_letter].width = c_width
        
    # Automatyczne dopasowanie szerokości kolumny B w Planie Treningowym
    max_len_plan_b = max((len(str(ws_plan.cell(row=r, column=2).value or '')) for r in range(1, row_idx)), default=30)
    ws_plan.column_dimensions['B'].width = max(30, max_len_plan_b + 3)

    # --- ZAKŁADKA 2: DZIENNIK TRENINGOWY (Cały Rok - 52 tyg) ---
    ws_dz = wb.create_sheet("DZIENNIK TRENINGOWY")
    for w in range(52):
        start_col = 3 + w*18
        ws_dz.merge_cells(start_row=1, start_column=start_col, end_row=1, end_column=start_col+17)
        c = ws_dz.cell(row=1, column=start_col, value=f"Tydzień {w+1}")
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
        c.font = Font(bold=True)
        for s in range(6):
            s_col = start_col + s*3
            ws_dz.merge_cells(start_row=2, start_column=s_col, end_row=2, end_column=s_col+2)
            c2 = ws_dz.cell(row=2, column=s_col, value=f"SERIA {s+1}")
            c2.alignment = Alignment(horizontal="center", vertical="center")
            c2.fill = PatternFill(start_color="E6F2FF", end_color="E6F2FF", fill_type="solid")
            c2.font = Font(bold=True)
            for idx_h, val_h in enumerate(["CIĘŻAR", "POWT.", "RIR"]):
                c3 = ws_dz.cell(row=3, column=s_col+idx_h, value=val_h)
                c3.alignment = Alignment(horizontal="center")
                c3.font = Font(bold=True)
                ws_dz.column_dimensions[get_column_letter(s_col+idx_h)].width = 9

    ws_dz.column_dimensions['A'].width = 12
    ws_dz.cell(row=3, column=1, value="KOLEJNOŚĆ").font = Font(bold=True)
    ws_dz.cell(row=3, column=2, value="ĆWICZENIE").font = Font(bold=True)
    
    r_dz = 4; ex_nr = 1; dzien_idx = 1; analiza_data = []
    for kat, cw in st.session_state.wylosowany_plan_cache:
        if kat == "NAGŁÓWEK DNIA":
            if r_dz > 4: r_dz += 1
            ws_dz.merge_cells(start_row=r_dz, start_column=1, end_row=r_dz, end_column=2)
            c = ws_dz.cell(row=r_dz, column=1, value=f"DZIEŃ {dzien_idx} - {cw['nazwa']}")
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = PatternFill(start_color="007BFF", end_color="007BFF", fill_type="solid")
            c.alignment = Alignment(horizontal="left", vertical="center")
            r_dz += 1; ex_nr = 1; dzien_idx += 1
        else:
            ws_dz.cell(row=r_dz, column=1, value=ex_nr).alignment = Alignment(horizontal="center")
            ws_dz.cell(row=r_dz, column=2, value=cw['nazwa'])
            analiza_data.append((cw['nazwa'], r_dz)); r_dz += 1; ex_nr += 1

    # Rysowanie krawędzi, ustawianie grubych linii oraz automatycznego zawijania tekstu dla kolumny B
    max_col_dz = 2 + 52 * 18
    ostatni_wiersz_siatki = max(100, r_dz + 15)
    
    for row in range(1, ostatni_wiersz_siatki + 1):
        for col in range(1, max_col_dz + 1):
            cell = ws_dz.cell(row=row, column=col)
            
            b_left = thin_side
            b_right = thin_side
            b_top = thin_side
            b_bottom = thin_side
            
            if col == 2: b_right = medium_side
            if col == 3: b_left = medium_side
            
            if col > 2:
                if (col - 2) % 18 == 1: b_left = medium_side
                if (col - 2) % 18 == 0: b_right = medium_side
                
            if row == 3:
                b_bottom = medium_side
                
            cell.border = Border(left=b_left, right=b_right, top=b_top, bottom=b_bottom)
            
            # NOWOŚĆ: Automatyczne włączenie zawijania tekstu dla komórek z nazwami ćwiczeń
            if col == 2 and row >= 4:
                cell.alignment = Alignment(wrap_text=True, vertical='center', horizontal='left')

    # NOWOŚĆ: Automatyczne dopasowanie szerokości kolumny B na podstawie najdłuższego ćwiczenia
    max_len_dz_b = max((len(str(ws_dz.cell(row=r, column=2).value or '')) for r in range(4, r_dz)), default=30)
    ws_dz.column_dimensions['B'].width = max(30, max_len_dz_b + 3)

    # --- ZAKŁADKA 3: ANALIZA PROGRESU (Formuły makrocyklowe) ---
    ws_an = wb.create_sheet("ANALIZA")
    def set_block(start_col, title, cols_titles, fill_color):
        ws_an.merge_cells(start_row=1, start_column=start_col, end_row=1, end_column=start_col+2)
        c = ws_an.cell(row=1, column=start_col, value=title)
        c.alignment = Alignment(horizontal="center"); c.font = Font(bold=True); c.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid"); c.border = thin_border
        for i, t in enumerate(cols_titles):
            c2 = ws_an.cell(row=2, column=start_col+i, value=t)
            c2.alignment = Alignment(horizontal="center"); c2.font = Font(bold=True, size=10); c2.border = thin_border; ws_an.column_dimensions[get_column_letter(start_col+i)].width = 14

    ws_an.column_dimensions['A'].width = 30; ws_an.merge_cells("A1:A2")
    c = ws_an.cell(row=1, column=1, value="ĆWICZENIE")
    c.alignment = Alignment(horizontal="center", vertical="center"); c.fill = fill_h; c.font = font_w; c.border = thin_border
    
    for w in range(52): set_block(2 + w*3, f"Tydzień {w+1}", ["Max Ciężar", "Suma Powt.", "Zrealiz. Serie"], "E6F2FF")
    
    c_off = 2 + 52*3
    mw = [list(range(0,4)), list(range(4,8)), list(range(8,13)), list(range(13,17)), list(range(17,21)), list(range(21,26)), list(range(26,30)), list(range(30,34)), list(range(34,39)), list(range(39,43)), list(range(43,47)), list(range(47,52))]
    for m in range(12): set_block(c_off + m*3, f"Miesiąc {m+1}", ["Śr. Max Ciężar", "Suma Powt.", "Suma Serii"], "CCE5FF")
    c_off += 36; 
    for q in range(4): set_block(c_off + q*3, f"Kwartał {q+1}", ["Śr. Max Ciężar", "Suma Powt.", "Suma Serii"], "99CCFF")
    c_off += 12; 
    for h in range(2): set_block(c_off + h*3, f"Półrocze {h+1}", ["Śr. Max Ciężar", "Suma Powt.", "Suma Serii"], "66B2FF")
    c_off += 6; set_block(c_off, "Cały Rok", ["Śr. Max Ciężar", "Suma Powt.", "Suma Serii"], "3399FF")

    r_an = 3
    for nazwa, r_in_dz in analiza_data:
        ws_an.cell(row=r_an, column=1, value=nazwa).border = thin_border
        for w in range(52):
            w_cells = [f"'DZIENNIK TRENINGOWY'!{get_column_letter(3 + w*18 + s*3)}{r_in_dz}" for s in range(6)]
            p_cells = [f"'DZIENNIK TRENINGOWY'!{get_column_letter(4 + w*18 + s*3)}{r_in_dz}" for s in range(6)]
            ws_an.cell(row=r_an, column=2+w*3, value=f'=IF(COUNT({",".join(w_cells)})>0, MAX({",".join(w_cells)}), "")').border = thin_border
            ws_an.cell(row=r_an, column=2+w*3+1, value=f'=IF(SUM({",".join(p_cells)})>0, SUM({",".join(p_cells)}), "")').border = thin_border
            ws_an.cell(row=r_an, column=2+w*3+2, value=f'=IF(COUNT({",".join(w_cells)})>0, COUNT({",".join(w_cells)}), "")').border = thin_border
        
        def app_agg(cb, weeks):
            mx_c = [f"{get_column_letter(2 + w*3)}{r_an}" for w in weeks]
            pt_c = [f"{get_column_letter(2 + w*3 + 1)}{r_an}" for w in weeks]
            sr_c = [f"{get_column_letter(2 + w*3 + 2)}{r_an}" for w in weeks]
            ws_an.cell(row=r_an, column=cb, value=f'=IF(COUNT({",".join(mx_c)})>0, ROUND(AVERAGE({",".join(mx_c)}), 1), "")').border = thin_border
            ws_an.cell(row=r_an, column=cb+1, value=f'=IF(SUM({",".join(pt_c)})>0, SUM({",".join(pt_c)}), "")').border = thin_border
            ws_an.cell(row=r_an, column=cb+2, value=f'=IF(SUM({",".join(sr_c)})>0, SUM({",".join(sr_c)}), "")').border = thin_border
            
        c_off = 158
        for m in range(12): app_agg(c_off + m*3, mw[m])
        c_off += 36
        for q in range(4): app_agg(c_off + q*3, [w for m in range(q*3, q*3+3) for w in mw[m]])
        c_off += 12
        app_agg(c_off, [w for m in range(0, 6) for w in mw[m]]); app_agg(c_off + 3, [w for m in range(6, 12) for w in mw[m]])
        c_off += 6; app_agg(c_off, list(range(52)))
        r_an += 1

    bio = io.BytesIO(); wb.save(bio); return bio.getvalue()
# ==============================================================================
# UI - INTERFEJS STRONY WEBOWEJ
# ==============================================================================
st.markdown('<div id="poczatek-strony"></div>', unsafe_allow_html=True)
st.title("Fizjo Workout Ultimate K.T.")
st.markdown("Zintegrowane środowisko projektowania programów treningowych.")

# Kod stylizujący pływający przycisk "Wróć na górę"
# Kod stylizujący pływający przycisk "Wróć na górę" (Lewa strona)
# Kod stylizujący pływający przycisk "Wróć na górę" (Lewa strona, poza paskiem bocznym)
st.markdown("""
<style>
.back-to-top {
    position: fixed;
    bottom: 20px;
    left: 320px; /* Przesunięto o 320px w prawo, aby ominąć pasek boczny */
    background-color: #28a745;
    color: white;
    border-radius: 50%;
    width: 50px;
    height: 50px;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 24px;
    text-decoration: none;
    box-shadow: 0 4px 6px rgba(0,0,0,0.3);
    z-index: 9999;
    transition: 0.3s;
}
.back-to-top:hover {
    background-color: #218838;
    color: white !important;
}
</style>
<a href="#poczatek-strony" class="back-to-top" title="Wróć na górę">⬆️</a>
""", unsafe_allow_html=True)

with st.sidebar:
    st.header("🔑 Dostęp do AI")
    user_api_key = st.text_input("Twój klucz API Groq:", type="password", help="Pobierz darmowy klucz ze strony console.groq.com")
    st.caption("Nie masz klucza? Pobierz go bezpłatnie ze strony: [console.groq.com](https://console.groq.com/)")
    groq_client = Groq(api_key=user_api_key) if user_api_key else None

    st.divider()
    st.header("⚙️ Konfiguracja")
    
    st.markdown("""
    <style>
    div[role="radiogroup"] label {
        padding: 6px 10px;
        border-radius: 8px;
        margin-bottom: 4px;
        transition: 0.3s;
    }
    div[role="radiogroup"] label:has(div:contains("FIZJO")) {
        background-color: rgba(40, 167, 69, 0.15);
        border-left: 4px solid #28a745;
    }
    div[role="radiogroup"] label:has(div:contains("GYM")) {
        background-color: rgba(0, 123, 255, 0.15);
        border-left: 4px solid #007bff;
    }
    </style>
    """, unsafe_allow_html=True)
    
    profil = st.radio("Profil Silnika:", [
        "🟩 FIZJO: Kompleksowy (Wszystkie partie)", 
        "🟩 FIZJO: Automatyczny Split (Dni Tygodnia)",
        "🟩 FIZJO: Ukierunkowany - Tylko Oddechowe",
        "🟩 FIZJO: Ukierunkowany - Tylko Głowa/Szyja",
        "🟩 FIZJO: Ukierunkowany - Tylko Kończyna górna",
        "🟩 FIZJO: Ukierunkowany - Tylko Core (Tułów)",
        "🟩 FIZJO: Ukierunkowany - Tylko Kończyna dolna",
        "🟦 GYM: Ogólnorozwojowy (FBW - Całe Ciało)",
        "🟦 GYM: Automatyczny Split (Dni Tygodnia)",
        "🟦 GYM: Ukierunkowany - Klatka piersiowa",
        "🟦 GYM: Ukierunkowany - Plecy",
        "🟦 GYM: Ukierunkowany - Barki",
        "🟦 GYM: Ukierunkowany - Ręce",
        "🟦 GYM: Ukierunkowany - Nogi",
        "🟦 GYM: Ukierunkowany - Pośladki",
        "🟦 GYM: Ukierunkowany - Brzuch / Core"
    ])
    
    is_gym = "GYM:" in profil
    czy_split = "Split" in profil
    czy_wielo_dniowy = czy_split or "Kompleksowy" in profil
    
    budzet = st.number_input("Ilość ćw. NA PARTIĘ:" if is_gym else "Budżet czasu (min):", min_value=1, max_value=120, value=4 if is_gym else 45)
    dni = st.number_input("Liczba dni (tylko dla Split/Kompleksowy):", min_value=1, max_value=31, value=4, disabled=not czy_wielo_dniowy)
    
    if st.button("⚡ GENERUJ AUTOMAT", use_container_width=True, type="primary"):
        generuj_plan(profil, budzet, dni)
        st.rerun()
        
    if st.button("❌ CZYŚĆ EKRAN (RESET)", use_container_width=True):
        st.session_state.wylosowany_plan_cache = []
        st.rerun()    
        
    st.divider()
    if st.session_state.wylosowany_plan_cache:
        st.success("Plan gotowy!")
        st.download_button("💾 POBIERZ RAPORT DOCX", generuj_docx(), "Raport_Pacjenta.docx", use_container_width=True)
        
        # --- ULTIMATE DETEKTOR TYPU PLANU ---
        czy_split = False
        slowa_kluczowe = ["split", "spit", "dni tygodnia", "poniedziałek", "wtorek", "środa", "czwartek", "piątek", "sobota", "niedziela"]
        
        for element in st.session_state.wylosowany_plan_cache:
            if any(slowo in str(element).lower() for slowo in slowa_kluczowe):
                czy_split = True
                break
                
        if not czy_split:
            for klucz, wartosc in st.session_state.items():
                if klucz != "wylosowany_plan_cache" and any(slowo in str(wartosc).lower() for slowo in slowa_kluczowe):
                    czy_split = True
                    break
                        
        # --- RYSOWANIE PRZYCISKÓW Z WYKRYWANIEM AWARII ---
        import traceback
        
        if czy_split:
            try:
                dane_gym = generuj_excel_gym(dni)
                if dane_gym:
                    st.download_button("🏋️‍♂️ GENERUJ EXCEL (DZIENNIK 52 TYG)", dane_gym, "Dziennik_Treningowy_Split.xlsx", 
                                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                                       use_container_width=True, type="primary")
            except Exception as e:
                st.error(f"Błąd Excel GYM: {e}")
                st.code(traceback.format_exc())
        else:
            try:
                dane_fizjo = generuj_excel_fizjo(dni)
                if dane_fizjo:
                    st.download_button("📊 GENERUJ EXCEL (STANDARDOWY)", dane_fizjo, "Plan_Standardowy.xlsx", 
                                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                                       use_container_width=True, type="primary")
            except Exception as e:
                st.error(f"Błąd Excel FIZJO: {e}")
                st.code(traceback.format_exc())
# ZAKŁADKI GŁÓWNE
tab1, tab2, tab_protokoly, tab3, tab4 = st.tabs(["📝 Twój Plan", "➕ Kreator", "🏥 Protokoły Kliniczne", "✨ Asystent AI", "⚙️ Baza Ćwiczeń"])

# ZAKŁADKA 1: WYGENEROWANY PLAN
# ZAKŁADKA 1: WYGENEROWANY PLAN
with tab1:
    if not st.session_state.wylosowany_plan_cache:
        st.info("👈 Użyj panelu bocznego, aby wygenerować plan lub przejdź do Kreatora.")
    else:
        def odswiez_nazwy_dni(blocks):
            dni_tyg = ["Poniedziałek", "Wtorek", "Środa", "Czwartek", "Piątek", "Sobota", "Niedziela"]
            for compliance_i, block in enumerate(blocks):
                kat, cw = block[0]
                if kat == "NAGŁÓWEK DNIA":
                    typ = cw.get("typ", "")
                    if typ == "Split":
                        nowy_dzien = dni_tyg[compliance_i] if compliance_i < 7 else f"Dzień {compliance_i+1}"
                        cw["nazwa"] = f"{nowy_dzien}: {cw['partie']}"
                    elif typ == "Kompleksowy":
                        cw["nazwa"] = f"Dzień {compliance_i+1} (FIZJO: Kompleksowy)"

        days_blocks = []
        current_block = []
        for item in st.session_state.wylosowany_plan_cache:
            if item[0] == "NAGŁÓWEK DNIA":
                if current_block: days_blocks.append(current_block)
                current_block = [item]
            else: current_block.append(item)
        if current_block: days_blocks.append(current_block)

        abs_idx = 0
        licznik = 1
        
        for block_idx, block in enumerate(days_blocks):
            for item_idx_in_block, (kat, cw) in enumerate(block):
                if kat == "NAGŁÓWEK DNIA":
                    st.markdown(f"### 📅 {cw['nazwa']}")
                    c1, c2, c3, c4 = st.columns([1.5, 1.5, 2, 5])
                    if block_idx > 0:
                        if c1.button("⬆️ Dzień", key=f"up_day_{abs_idx}", use_container_width=True):
                            days_blocks[block_idx], days_blocks[block_idx-1] = days_blocks[block_idx-1], days_blocks[block_idx]
                            odswiez_nazwy_dni(days_blocks)
                            st.session_state.wylosowany_plan_cache = [i for b in days_blocks for i in b]
                            st.rerun()
                    if block_idx < len(days_blocks) - 1:
                        if c2.button("⬇️ Dzień", key=f"down_day_{abs_idx}", use_container_width=True):
                            days_blocks[block_idx], days_blocks[block_idx+1] = days_blocks[block_idx+1], days_blocks[block_idx]
                            odswiez_nazwy_dni(days_blocks)
                            st.session_state.wylosowany_plan_cache = [i for b in days_blocks for i in b]
                            st.rerun()
                    if c3.button("❌ Usuń cały", key=f"del_day_{abs_idx}", type="primary", use_container_width=True):
                        days_blocks.pop(block_idx)
                        odswiez_nazwy_dni(days_blocks)
                        st.session_state.wylosowany_plan_cache = [i for b in days_blocks for i in b]
                        st.rerun()
                    licznik = 1
                    abs_idx += 1
                    continue
                    
                # B. Pojedyncze ćwiczenie
                with st.expander(f"{licznik}. {cw['nazwa']} ({kat})", expanded=True):
                    col1, col2 = st.columns([3, 2])
                    with col1:
                        nowe_parametry = st.text_input("Zalecenie", cw['parametry'], key=f"p_{abs_idx}")
                        nowe_uwagi = st.text_input("Uwagi", cw.get('uwagi', ''), key=f"u_{abs_idx}")
                        st.session_state.wylosowany_plan_cache[abs_idx][1]['parametry'] = nowe_parametry
                        st.session_state.wylosowany_plan_cache[abs_idx][1]['uwagi'] = nowe_uwagi
                        st.caption(f"**Anatomia:** {cw['miesnie']}")
                        st.write(cw['opis'])
                    with col2:
                        c_up, c_down, c_swap, c_del = st.columns(4)
                        moze_w_gore = abs_idx > 0 and st.session_state.wylosowany_plan_cache[abs_idx-1][0] != "NAGŁÓWEK DNIA"
                        moze_w_dol = abs_idx < len(st.session_state.wylosowany_plan_cache) - 1 and st.session_state.wylosowany_plan_cache[abs_idx+1][0] != "NAGŁÓWEK DNIA"
                        
                        if moze_w_gore and c_up.button("⬆️", key=f"up_{abs_idx}"):
                            st.session_state.wylosowany_plan_cache[abs_idx], st.session_state.wylosowany_plan_cache[abs_idx-1] = st.session_state.wylosowany_plan_cache[abs_idx-1], st.session_state.wylosowany_plan_cache[abs_idx]
                            st.rerun()
                            
                        if moze_w_dol and c_down.button("⬇️", key=f"down_{abs_idx}"):
                            st.session_state.wylosowany_plan_cache[abs_idx], st.session_state.wylosowany_plan_cache[abs_idx+1] = st.session_state.wylosowany_plan_cache[abs_idx+1], st.session_state.wylosowany_plan_cache[abs_idx]
                            st.rerun()
                            
                        if c_swap.button("🔄", key=f"swap_{abs_idx}", help="Wylosuj i wstaw inne ćwiczenie z tej samej partii"):
                            czysta_kat = kat.replace("GYM: ", "").replace(" (Rozgrzewka)", "").replace(" (Wyciszenie)", "")
                            if kat == "GYM: Zakończenie":
                                czysta_kat = "Zakończenie treningu"
                                
                            baza_docelowa = GLOBALNA_BAZA.get(czysta_kat, [])
                            if baza_docelowa:
                                nowe_cw = random.choice(baza_docelowa).copy()
                                nowe_cw["uwagi"] = ""
                                st.session_state.wylosowany_plan_cache[abs_idx] = (kat, nowe_cw)
                                st.rerun()
                                
                        if c_del.button("❌", key=f"del_{abs_idx}", type="primary"):
                            st.session_state.wylosowany_plan_cache.pop(abs_idx)
                            st.rerun()
                            
                licznik += 1
                abs_idx += 1

# ZAKŁADKA 2: KREATOR MANUALNY
with tab2:
    st.subheader("Manualne dodawanie ćwiczeń")
    podzak_fizjo, podzak_gym = st.tabs(["🏥 Baza Fizjoterapia", "🏋️ Baza Siłownia"])
    
    with podzak_fizjo:
        szukaj_f = st.text_input("🔍 Wyszukaj ćwiczenie z Fizjoterapii:", key="sz_f").strip().lower()
        st.divider()
        if szukaj_f:
            for kat, lista in BAZA_FIZJO.items():
                for cw in lista:
                    if szukaj_f in cw['nazwa'].lower() or szukaj_f in cw['miesnie'].lower():
                        c1, c2 = st.columns([4, 1])
                        c1.write(f"**{cw['nazwa']}** ({kat})\n↳ [{cw['miesnie']}]")
                        if c2.button("Dodaj", key=f"df_{cw['nazwa']}"):
                            st.session_state.wylosowany_plan_cache.append((kat, cw.copy()))
                            st.toast("Dodano!")
        else:
            kat_wyb_f = st.radio("Wybierz kategorię Fizjo:", list(BAZA_FIZJO.keys()))
            for cw in BAZA_FIZJO[kat_wyb_f]:
                c1, c2 = st.columns([4, 1])
                c1.write(f"**{cw['nazwa']}** \n↳ [{cw['miesnie']}]")
                if c2.button("Dodaj", key=f"bf_{cw['nazwa']}"):
                    st.session_state.wylosowany_plan_cache.append((kat_wyb_f, cw.copy()))
                    st.toast("Dodano!")

    with podzak_gym:
        szukaj_g = st.text_input("🔍 Wyszukaj ćwiczenie z Siłowni:", key="sz_g").strip().lower()
        st.divider()
        if szukaj_g:
            for kat, lista in BAZA_SILOWNIA.items():
                for cw in lista:
                    if szukaj_g in cw['nazwa'].lower() or szukaj_g in cw['miesnie'].lower():
                        c1, c2 = st.columns([4, 1])
                        c1.write(f"**{cw['nazwa']}** ({kat})\n↳ [{cw['miesnie']}]")
                        if c2.button("Dodaj", key=f"dg_{cw['nazwa']}"):
                            st.session_state.wylosowany_plan_cache.append((f"GYM: {kat}", cw.copy()))
                            st.toast("Dodano!")
        else:
            kat_wyb_g = st.radio("Wybierz kategorię Siłownia:", list(BAZA_SILOWNIA.keys()))
            for cw in BAZA_SILOWNIA[kat_wyb_g]:
                c1, c2 = st.columns([4, 1])
                c1.write(f"**{cw['nazwa']}** \n↳ [{cw['miesnie']}]")
                if c2.button("Dodaj", key=f"bg_{cw['nazwa']}"):
                    st.session_state.wylosowany_plan_cache.append((f"GYM: {kat_wyb_g}", cw.copy()))
                    st.toast("Dodano!")

# NOWA ZAKŁADKA: PROTOKOŁY KLINICZNE (PANEL GŁÓWNY)
with tab_protokoly:
    st.subheader("🏥 Gotowe Protokoły Kliniczne")
    st.markdown("Wyszukaj lub wybierz jednostkę z odpowiedniej kategorii. Protokoły są ułożone alfabetycznie.")
    
    if PROTOKOLY_KLINICZNE:
        szukana_choroba = st.text_input("🔍 Wyszukaj jednostkę chorobową (np. 'rwa', 'zespół'):", key="szukaj_proto").strip().lower()
        st.divider()
        
        znaleziono_protokol = False
        
        # Iteracja po głównych kategoriach (Kręgosłup, Kończyna itd.)
        for kategoria_chorob, choroby_w_kategorii in PROTOKOLY_KLINICZNE.items():
            
            # Filtrowanie i sortowanie alfabetyczne chorób w locie
            dopasowane = {choroba: cwiczenia for choroba, cwiczenia in choroby_w_kategorii.items() if szukana_choroba in choroba.lower()}
            dopasowane_posortowane = dict(sorted(dopasowane.items()))
            
            if dopasowane_posortowane:
                znaleziono_protokol = True
                st.markdown(f"#### 🔹 {kategoria_chorob}") # Elegancki nagłówek grupy
                
                for choroba, lista_cwiczen in dopasowane_posortowane.items():
                    with st.expander(f"⚕️ {choroba} (Podgląd: {len(lista_cwiczen)} ćw.)", expanded=False):
                        st.markdown("**Lista ćwiczeń w tym protokole:**")
                        for nazwa_cw in lista_cwiczen:
                            st.markdown(f"• {nazwa_cw}")
                        
                        st.write("") 
                        if st.button("🚀 URUCHOM TEN PROTOKÓŁ", key=f"proto_btn_{choroba}", use_container_width=True, type="primary"):
                            generuj_protokol(choroba)
                            st.rerun()
                st.write("") # Odstęp między sekcjami
                
        if not znaleziono_protokol:
            st.info("Nie znaleziono protokołu dla wpisanej frazy.")
    else:
        st.error("Brak pliku `protokoly.json` w katalogu aplikacji. Utwórz plik, aby aktywować bazę kliniczną.")
# ZAKŁADKA 3: CZAT AI GROQ
with tab3:
    st.subheader("Wirtualny Konsultant Treningowy")
    if not groq_client:
        st.info("👈 Wklej klucz w panelu bocznym.")
    else:
        for msg in st.session_state.historia_wiadomosci:
            if msg["role"] != "system":
                with st.chat_message(msg["role"]): st.markdown(msg["content"])
        if prompt := st.chat_input("Pytanie do AI..."):
            with st.chat_message("user"): st.markdown(prompt)
            st.session_state.historia_wiadomosci.append({"role": "user", "content": prompt})
            with st.chat_message("assistant"):
                try:
                    completion = groq_client.chat.completions.create(model="llama-3.3-70b-versatile", messages=st.session_state.historia_wiadomosci, temperature=0.7, max_tokens=1024)
                    odp = completion.choices[0].message.content
                    st.markdown(odp)
                    st.session_state.historia_wiadomosci.append({"role": "assistant", "content": odp})
                except Exception as e: st.error(f"Błąd: {e}")

# ZAKŁADKA 4: MENEDŻER WŁASNYCH ĆWICZEŃ
with tab4:
    st.subheader("Zarządzaj własną bazą")
    with st.expander("➕ Dodaj nowe ćwiczenie", expanded=False):
        with st.form("form_dodaj_cwiczenie", clear_on_submit=True):
            kategorie = [f"FIZJO: {k}" for k in BAZA_FIZJO.keys()] + [f"GYM: {k}" for k in BAZA_SILOWNIA.keys()]
            kat_wyb = st.selectbox("Kategoria Docelowa:", kategorie)
            n_nazwa = st.text_input("Nazwa ćwiczenia:")
            n_parametry = st.text_input("Zalecenie:")
            c1, c2 = st.columns(2)
            n_czas = c1.number_input("Czas (min):", min_value=1, value=2)
            n_miesnie = c2.text_input("Anatomia:")
            n_opis = st.text_area("Opis:")
            
            if st.form_submit_button("Zapisz", type="primary"):
                if n_nazwa and n_parametry and n_miesnie and n_opis:
                    nowe_cw = {"nazwa": n_nazwa, "opis": n_opis, "czas_min": int(n_czas), "parametry": n_parametry, "miesnie": n_miesnie}
                    tb, kategoria_str = kat_wyb.split(": ")
                    dane = {"FIZJO": {}, "GYM": {}}
                    if os.path.exists(PLIK_WLASNYCH_CWICZEN):
                        try:
                            with open(PLIK_WLASNYCH_CWICZEN, "r") as f: dane = json.load(f)
                        except: pass
                    if tb not in dane: dane[tb] = {}
                    if kategoria_str not in dane[tb]: dane[tb][kategoria_str] = []
                    dane[tb][kategoria_str].append(nowe_cw)
                    with open(PLIK_WLASNYCH_CWICZEN, "w") as f: json.dump(dane, f, ensure_ascii=False, indent=4)
                    st.success("Dodano! Odśwież stronę.")
